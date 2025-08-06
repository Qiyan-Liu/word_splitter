#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档拆分工具

作者: QiyanLiu
日期: 2025-08-06

根据文档目录结构将大型Word文档拆分为多个小文档
支持多线程处理，保持原始格式和样式
"""

import os
import sys
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import List, Dict, Tuple, Optional
import logging
from dataclasses import dataclass
from collections import defaultdict
import gc
from functools import lru_cache
import weakref

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("请安装python-docx库: pip install python-docx")
    sys.exit(1)

# 配置日志 - 只记录到文件，控制台输出由主程序控制
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('word_splitter.log', encoding='utf-8')
    ]
)
logger = logging.getLogger(__name__)

@dataclass
class ChapterInfo:
    """章节信息"""
    title: str
    level: int
    start_paragraph: int
    end_paragraph: int
    paragraphs: List[int]

class WordDocumentSplitter:
    """Word文档拆分器"""
    
    def __init__(self, min_level: int = 3, max_workers_docs: int = 4, max_workers_chapters: int = 2):
        """
        初始化文档拆分器
        
        Args:
            min_level: 最小拆分层级
            max_workers_docs: 处理多个文档的线程数
            max_workers_chapters: 处理单个文档章节的线程数
        """
        self.min_level = min_level
        self.max_workers_docs = max_workers_docs
        self.max_workers_chapters = max_workers_chapters
        self.lock = threading.Lock()
        
        # 移除缓存机制以避免哈希问题
    

    
    def analyze_document_structure(self, doc_path: str) -> Tuple[Document, List[ChapterInfo]]:
        """分析文档结构，识别章节（性能优化版）"""
        import time
        start_time = time.time()
        
        try:
            logger.info(f"开始分析文档结构: {doc_path}")
            

            
            doc = Document(doc_path)
            chapters = []
            current_levels = {}
            all_headings = []  # 存储所有标题信息
            
            # 预先获取所有段落（减少重复访问）
            paragraphs = doc.paragraphs
            total_paragraphs = len(paragraphs)
            
            logger.info(f"文档包含 {total_paragraphs} 个段落")
            
            # 第一步：收集所有标题信息（批量处理）
            for i, paragraph in enumerate(paragraphs):
                outline_level = self._get_outline_level(paragraph)
                text = paragraph.text.strip()
                
                # 只要有大纲级别且有文本内容，就认为是标题
                if outline_level > 0 and text:
                    all_headings.append({
                        'level': outline_level,
                        'title': text,
                        'paragraph_index': i
                    })
                
                # 每处理1000个段落记录一次进度
                if (i + 1) % 1000 == 0:
                    logger.debug(f"已处理 {i + 1}/{total_paragraphs} 个段落")
            
            logger.info(f"找到 {len(all_headings)} 个标题")
            
            # 第二步：分析每个分支的最深层级，确定拆分点
            if all_headings:
                # 第三步：遍历标题，在合适的位置创建章节
                for i, heading in enumerate(all_headings):
                    outline_level = heading['level']
                    paragraph_index = heading['paragraph_index']
                    
                    # 更新当前层级信息
                    current_levels[outline_level] = {
                        'title': heading['title'],
                        'paragraph_index': paragraph_index
                    }
                    
                    # 清除更深层级的信息
                    keys_to_remove = [k for k in current_levels.keys() if k > outline_level]
                    for k in keys_to_remove:
                        del current_levels[k]
                    
                    # 判断是否应该在此处创建章节
                    should_create = self._should_create_chapter_at_position(all_headings, i, outline_level)
                    
                    if should_create:
                        chapter_title = self._build_chapter_title(current_levels, outline_level)
                        chapter = ChapterInfo(
                            title=chapter_title,
                            level=outline_level,
                            start_paragraph=paragraph_index,
                            end_paragraph=paragraph_index,
                            paragraphs=[paragraph_index]
                        )
                        chapters.append(chapter)
            
            # 设置章节的结束段落
            self._set_chapter_boundaries(chapters, total_paragraphs, doc)
            
            # 性能统计
            end_time = time.time()
            processing_time = end_time - start_time
            
            logger.info(f"文档结构分析完成，共识别 {len(chapters)} 个章节，耗时 {processing_time:.2f} 秒")
            return doc, chapters
            
        except Exception as e:
            logger.error(f"分析文档结构失败: {e}")
            raise
    
    def _get_outline_level(self, paragraph) -> int:
        """获取段落的大纲级别"""
        return self._calculate_outline_level(paragraph)
    
    def _calculate_outline_level(self, paragraph) -> int:
        """计算段落的大纲级别（仅基于Word标准样式和大纲级别）"""
        try:
            # 1. 首先检查段落样式是否为标准标题样式
            style_name = paragraph.style.name
            
            # 检查英文标题样式 (Heading 1, Heading 2, etc.)
            if style_name.startswith('Heading '):
                level_str = style_name.replace('Heading ', '')
                try:
                    return int(level_str)
                except ValueError:
                    return 1  # 默认为1级标题
            
            # 检查中文标题样式 (标题 1, 标题 2, etc.)
            if '标题' in style_name:
                import re
                level_match = re.search(r'(\d+)', style_name)
                if level_match:
                    return int(level_match.group(1))
                return 1  # 默认为1级标题
            
            # 检查自定义样式（样式1, 样式2, 样式3等）
            if style_name.startswith('样式'):
                import re
                level_match = re.search(r'(\d+)', style_name)
                if level_match:
                    level_num = int(level_match.group(1))
                    # 样式3通常用作一级标题，样式4用作二级标题等
                    # 但为了保险起见，我们检查段落内容来判断
                    text = paragraph.text.strip()
                    if text and self._looks_like_heading(text):
                        return level_num
                return 0  # 不是标题样式
            
            # 2. 检查段落格式的大纲级别（Word内置的大纲级别属性）
            if hasattr(paragraph._element, 'pPr') and paragraph._element.pPr is not None:
                outline_lvl = paragraph._element.pPr.find(qn('w:outlineLvl'))
                if outline_lvl is not None:
                    level_value = int(outline_lvl.get(qn('w:val')))
                    return level_value + 1  # Word的大纲级别从0开始，我们从1开始
            
            # 3. 如果既不是标题样式，也没有大纲级别，则不是标题
            return 0
            
        except Exception:
            return 0
    
    def _looks_like_heading(self, text: str) -> bool:
        """判断文本是否看起来像标题"""
        if not text or len(text.strip()) == 0:
            return False
        
        text = text.strip()
        
        # 检查是否为章节标题格式（一、二、三、等）
        import re
        if re.match(r'^[一二三四五六七八九十]+、', text):
            return True
        
        # 检查是否为数字章节格式（1、2、3、等）
        if re.match(r'^\d+[、.]', text):
            return True
        
        # 检查是否为常见标题关键词
        heading_keywords = [
            '绪论', '引言', '前言', '概述', '背景', '意义', '目的',
            '文献综述', '理论基础', '研究方法', '分析', '讨论',
            '结论', '总结', '展望', '参考文献', '致谢', '谢辞', '附录'
        ]
        
        for keyword in heading_keywords:
            if keyword in text:
                return True
        
        # 检查长度（标题通常较短）
        if len(text) <= 50 and not '。' in text:
            # 短文本且不包含句号，可能是标题
            return True
        
        return False
    
    def _is_likely_toc_content(self, text: str) -> bool:
        """判断是否为目录内容"""
        # 目录特征：包含页码、点号连接、特定关键词等
        toc_indicators = [
            '目　　录',
            '目录',
            '......',
            '………',
            '.....',
            '-----',
            '参考文献',
            '致谢',
            '谢辞',
            '附录',
            '本科毕业论文',
            '专　　业',
            '院　　系',
            '指导教师',
            '摘　　要',
            'Abstract',
            'Keywords',
            '关键词',
            '曲式分析法',
            '音乐分析法',
            '文献分析法'
        ]
        
        # 检查是否包含目录特征
        for indicator in toc_indicators:
            if indicator in text:
                return True
        
        # 检查是否为页码格式（数字结尾）
        import re
        if re.search(r'\d+\s*$', text):
            return True
        
        # 检查是否包含大量空格（目录对齐格式）
        if text.count('　') > 2 or text.count(' ') > 10:
            return True
        
        # 检查是否为论文封面信息
        if any(keyword in text for keyword in ['专业：', '院系：', '指导教师：', '年　月']):
            return True
            
        return False
    
    def _is_likely_heading(self, paragraph, text: str) -> bool:
        """判断是否为真正的标题"""
        # 标题特征检查
        if not paragraph.runs:
            return False
        
        # 排除明显不是标题的内容
        import re
        
        # 排除表格相关内容（以冒号结尾的描述性文本）
        if re.search(r'(表格|图片|图表).*[：:]$', text):
            return False
        
        # 排除结束语句
        if re.search(r'(结束|完毕|完成)[。.]$', text):
            return False
        
        # 排除过长的文本（标题通常较短）
        if len(text) > 50:
            return False
        
        # 检查字体加粗
        is_bold = any(run.bold for run in paragraph.runs if run.bold is not None)
        
        # 检查字体大小
        font_size = paragraph.runs[0].font.size
        is_large_font = font_size and font_size.pt >= 12
        
        # 检查是否包含章节编号
        has_chapter_number = bool(re.search(r'^(第[一二三四五六七八九十\d]+[章节部分]|\d+[\.]\d*\s*|[一二三四五六七八九十]+[、\.]\s*)', text))
        
        # 检查是否为学术论文常见章节标题
        academic_headings = [
            '绪论', '引言', '前言', '概述',
            '研究背景', '研究现状', '研究意义', '研究目的', '研究方法',
            '文献综述', '理论基础', '相关工作',
            '实验设计', '实验结果', '结果分析',
            '讨论', '分析', '总结', '结论',
            '展望', '建议', '不足',
            '艺术特征', '音乐分析', '演奏技巧', '创作背景',
            '基础内容', '进阶内容', '基础概念', '实践应用', '应用场景', '最佳实践'
        ]
        
        # 检查是否包含学术标题关键词
        has_academic_keywords = any(
            keyword in text for keyword in academic_headings
        )
        
        # 综合判断
        score = 0
        if is_bold:
            score += 2
        if is_large_font:
            score += 1
        if has_chapter_number:
            score += 3  # 章节编号权重更高
        if has_academic_keywords:
            score += 2
        
        # 降低阈值，分数>=3认为是标题（平衡严格性和识别率）
        return score >= 3
    
    def _should_create_chapter_at_position(self, all_headings: List[Dict], current_index: int, current_level: int) -> bool:
        """判断是否应该在当前位置创建章节（局部适应章节层级）"""
        # 分析当前章节分支的最大层级
        branch_max_level = self._get_branch_max_level(all_headings, current_index, current_level)
        
        # 如果当前分支的最大层级小于min_level，则使用分支的最大层级作为拆分层级
        effective_min_level = min(self.min_level, branch_max_level)
        
        # 只在达到有效的最小层级时创建章节
        if current_level == effective_min_level:
            return True
        
        return False
    
    def _get_branch_max_level(self, all_headings: List[Dict], current_index: int, current_level: int) -> int:
        """获取当前章节分支的最大层级"""
        # 构建当前章节的路径（从根到当前位置）
        current_path = {}
        for i in range(current_index + 1):
            heading = all_headings[i]
            if heading['level'] <= current_level:
                current_path[heading['level']] = heading['title']
            if heading['level'] == current_level and i == current_index:
                break
        
        # 查找属于当前分支的所有后续标题，并找出最大层级
        max_level = current_level
        
        for i in range(current_index + 1, len(all_headings)):
            heading = all_headings[i]
            
            # 如果遇到同级或更高级的标题，说明当前分支结束
            if heading['level'] <= current_level:
                break
            
            # 检查是否属于当前分支
            if self._belongs_to_current_branch(all_headings, i, current_path, current_level):
                max_level = max(max_level, heading['level'])
        
        return max_level
    
    def _belongs_to_current_branch(self, all_headings: List[Dict], heading_index: int, current_path: Dict, current_level: int) -> bool:
        """判断指定的标题是否属于当前分支"""
        # 重建到指定标题的路径
        path_to_heading = {}
        
        for i in range(heading_index + 1):
            heading = all_headings[i]
            if heading['level'] <= current_level:
                path_to_heading[heading['level']] = heading['title']
        
        # 检查路径是否匹配
        for level in range(1, current_level + 1):
            if level in current_path and level in path_to_heading:
                if current_path[level] != path_to_heading[level]:
                    return False
            elif level in current_path or level in path_to_heading:
                return False
        
        return True
    

    
    def _has_deeper_branch(self, all_headings: List[Dict], current_index: int, current_level: int) -> bool:
        """检查当前分支是否还有更深的层级"""
        # 构建当前路径
        current_path = {}
        for i in range(current_index + 1):
            heading = all_headings[i]
            if heading['level'] <= current_level:
                current_path[heading['level']] = heading['title']
            if heading['level'] == current_level and i == current_index:
                break
        
        # 查看后续是否有属于同一分支的更深层级
        for i in range(current_index + 1, len(all_headings)):
            future_heading = all_headings[i]
            future_level = future_heading['level']
            
            # 如果遇到同级或更浅的层级，说明当前分支结束
            if future_level <= current_level:
                break
            
            # 如果是更深层级，检查是否属于当前分支
            if future_level > current_level:
                # 构建到这个更深层级的路径
                future_path = {}
                for j in range(current_index + 1):
                    heading = all_headings[j]
                    if heading['level'] <= current_level:
                        future_path[heading['level']] = heading['title']
                
                # 检查路径是否匹配（即是否属于同一分支）
                path_matches = True
                for level in current_path:
                    if level not in future_path or current_path[level] != future_path[level]:
                        path_matches = False
                        break
                
                if path_matches:
                    return True
        
        return False
    
    def _determine_target_level(self, unique_levels: List[int]) -> int:
        """根据文档的整体结构确定目标拆分层级（保留用于向后兼容）"""
        if not unique_levels:
            return self.min_level
        
        # 检查层级是否连续
        is_continuous = True
        if len(unique_levels) > 1:
            for i in range(len(unique_levels) - 1):
                if unique_levels[i + 1] - unique_levels[i] > 1:
                    is_continuous = False
                    break
        
        max_level = max(unique_levels)
        
        if not is_continuous:
            # 层级不连续，找到小于min_level的最深层级
            valid_levels = [level for level in unique_levels if level < self.min_level]
            if valid_levels:
                return max(valid_levels)
            else:
                # 如果没有小于min_level的层级，返回最小层级
                return min(unique_levels)
        
        # 层级连续的情况
        # 如果最大层级已达到或超过min_level，在min_level创建章节
        if max_level >= self.min_level:
            return self.min_level
        
        # 如果最大层级小于min_level，在最深层级创建章节（但至少是第2层）
        return max_level if max_level >= 2 else 2
    
    def _should_create_chapter(self, current_levels: Dict, outline_level: int) -> bool:
        """判断是否应该创建新章节"""
        # 获取当前存在的层级
        existing_levels = sorted(current_levels.keys())
        
        if not existing_levels:
            return False
        
        # 检查层级是否连续
        is_continuous = True
        if len(existing_levels) > 1:
            for i in range(len(existing_levels) - 1):
                if existing_levels[i + 1] - existing_levels[i] > 1:
                    is_continuous = False
                    break
        
        max_level = max(existing_levels)
        
        if not is_continuous:
            # 层级不连续，找到小于min_level的最深层级
            valid_levels = [level for level in existing_levels if level < self.min_level]
            if valid_levels:
                target_level = max(valid_levels)
                return outline_level == target_level and target_level >= 2
            else:
                # 如果没有小于min_level的层级，则不创建章节
                return False
        
        # 层级连续的情况
        # 如果最大层级已达到或超过min_level，只在min_level创建章节
        if max_level >= self.min_level:
            return outline_level == self.min_level
        
        # 如果最大层级小于min_level，在最深层级创建章节
        return outline_level == max_level and max_level >= 2
    
    def _build_chapter_title(self, current_levels: Dict, target_level: int) -> str:
        """构建章节标题"""
        titles = []
        for level in sorted(current_levels.keys()):
            if level <= target_level:
                titles.append(current_levels[level]['title'])
        return ' - '.join(titles) if titles else f"章节_{target_level}"
    
    def _set_chapter_boundaries(self, chapters: List[ChapterInfo], total_paragraphs: int, doc: Document = None):
        """设置章节边界（智能处理层级跳跃）"""
        if not chapters:
            return
        
        # 获取文档中所有标题的位置信息
        all_heading_positions = []
        
        if doc:
            for i, paragraph in enumerate(doc.paragraphs):
                outline_level = self._get_outline_level(paragraph)
                text = paragraph.text.strip()
                if outline_level > 0 and text:
                    all_heading_positions.append({
                        'paragraph_index': i,
                        'level': outline_level,
                        'title': text
                    })
        
        for i, chapter in enumerate(chapters):
            if i < len(chapters) - 1:
                next_chapter = chapters[i + 1]
                
                # 智能确定章节结束位置
                # 查找当前章节和下一章节之间是否有同级或更高级的标题
                chapter_end = next_chapter.start_paragraph - 1
                
                # 检查是否有层级跳跃导致的内容包含问题
                for heading_info in all_heading_positions:
                    heading_pos = heading_info['paragraph_index']
                    heading_level = heading_info['level']
                    
                    # 如果在当前章节范围内发现了同级或更高级的标题
                    if (chapter.start_paragraph < heading_pos < next_chapter.start_paragraph and 
                        heading_level <= chapter.level):
                        # 将章节结束位置调整到该标题之前
                        chapter_end = heading_pos - 1
                        break
                
                chapter.end_paragraph = max(chapter.start_paragraph, chapter_end)
            else:
                # 最后一个章节
                chapter.end_paragraph = total_paragraphs - 1
            
            # 设置章节包含的段落
            chapter.paragraphs = list(range(chapter.start_paragraph, chapter.end_paragraph + 1))
    
    def create_chapter_document(self, original_doc: Document, chapter: ChapterInfo, output_path: str):
        """创建章节文档（性能优化版）"""
        try:
            # 创建新文档
            new_doc = Document()
            
            # 复制样式
            self._copy_styles(original_doc, new_doc)
            
            # 批量复制章节内容（减少单次调用开销）
            self._copy_paragraphs_batch(original_doc, new_doc, chapter.paragraphs)
            
            # 复制表格（如果在章节范围内）
            self._copy_tables_in_range(original_doc, new_doc, chapter.paragraphs)
            
            # 保存文档
            safe_filename = self._sanitize_filename(chapter.title)
            full_output_path = os.path.join(output_path, f"{safe_filename}.docx")
            
            with self.lock:
                os.makedirs(os.path.dirname(full_output_path), exist_ok=True)
                new_doc.save(full_output_path)
            
            # 强制垃圾回收以释放内存
            del new_doc
            gc.collect()
            
            logger.info(f"章节文档已创建: {full_output_path}")
            return full_output_path
            
        except Exception as e:
            logger.error(f"创建章节文档失败: {e}")
            raise
    
    def _copy_paragraphs_batch(self, source_doc: Document, target_doc: Document, paragraph_indices: List[int]):
        """批量复制段落（性能优化）"""
        try:
            source_paragraphs = source_doc.paragraphs
            source_part = source_doc.part
            
            # 批量处理段落，减少函数调用开销
            for para_index in paragraph_indices:
                if para_index < len(source_paragraphs):
                    original_para = source_paragraphs[para_index]
                    self._copy_paragraph(original_para, target_doc, source_part)
                    
        except Exception as e:
            logger.warning(f"批量复制段落失败: {e}")
            # 回退到逐个复制
            for para_index in paragraph_indices:
                if para_index < len(source_doc.paragraphs):
                    try:
                        original_para = source_doc.paragraphs[para_index]
                        self._copy_paragraph(original_para, target_doc, source_doc.part)
                    except Exception as inner_e:
                        logger.warning(f"复制段落 {para_index} 失败: {inner_e}")
                        continue
    
    def _copy_styles(self, source_doc: Document, target_doc: Document):
        """复制文档样式"""
        try:
            # 获取目标文档现有样式名称（一次性获取，避免重复计算）
            existing_styles = {s.name for s in target_doc.styles}
            
            # 复制段落样式
            for style in source_doc.styles:
                if style.name not in existing_styles:
                    try:
                        target_doc.styles.add_style(style.name, style.type)
                    except Exception:
                        pass  # 忽略样式复制错误
            
        except Exception as e:
            logger.warning(f"样式复制警告: {e}")
    
    def _copy_paragraph(self, source_para, target_doc: Document, source_doc_part):
        """复制段落，包括文本、格式和图片"""
        try:
            new_para = target_doc.add_paragraph()
            
            # 安全复制段落格式
            try:
                if hasattr(source_para, 'style') and source_para.style:
                    style_name = str(source_para.style.name) if source_para.style.name else None
                    if style_name:
                        # 检查目标文档是否有此样式
                        for style in target_doc.styles:
                            if style.name == style_name:
                                new_para.style = style
                                break
            except Exception as e:
                logger.debug(f"样式复制失败: {e}")
            
            # 安全复制段落对齐
            try:
                if hasattr(source_para, 'alignment') and source_para.alignment is not None:
                    new_para.alignment = source_para.alignment
            except Exception as e:
                logger.debug(f"对齐复制失败: {e}")
            
            # 复制runs
            for run in source_para.runs:
                try:
                    # 检查run中是否包含图片
                    images_in_run = self._get_images_from_run(run, source_doc_part)
                    
                    if images_in_run:
                        # 如果run包含图片，先添加文本，然后添加图片
                        if run.text:
                            new_run = new_para.add_run(run.text)
                            self._copy_run_format(run, new_run)
                        
                        # 添加图片
                        for image_data in images_in_run:
                            try:
                                new_run = new_para.add_run()
                                new_run.add_picture(image_data['stream'], width=image_data['width'], height=image_data['height'])
                            except Exception as e:
                                logger.warning(f"图片复制失败: {e}")
                                # 如果图片复制失败，添加占位文本
                                new_run = new_para.add_run("[图片]")
                    else:
                        # 普通文本run
                        if run.text:  # 只复制有文本的run
                            new_run = new_para.add_run(run.text)
                            self._copy_run_format(run, new_run)
                except Exception as e:
                    logger.debug(f"Run复制失败: {e}")
                    continue
                
        except Exception as e:
            logger.warning(f"段落复制警告: {e}")
    
    def _copy_run_format(self, source_run, target_run):
        """安全复制run的格式"""
        try:
            # 安全复制字体格式
            if hasattr(source_run, 'font') and source_run.font:
                try:
                    if hasattr(source_run.font, 'name') and source_run.font.name:
                        target_run.font.name = str(source_run.font.name)
                except Exception:
                    pass
                
                try:
                    if hasattr(source_run.font, 'size') and source_run.font.size:
                        target_run.font.size = source_run.font.size
                except Exception:
                    pass
                
                try:
                    if hasattr(source_run.font, 'bold') and source_run.font.bold is not None:
                        target_run.font.bold = bool(source_run.font.bold)
                except Exception:
                    pass
                
                try:
                    if hasattr(source_run.font, 'italic') and source_run.font.italic is not None:
                        target_run.font.italic = bool(source_run.font.italic)
                except Exception:
                    pass
                
                try:
                    if hasattr(source_run.font, 'underline') and source_run.font.underline is not None:
                        target_run.font.underline = source_run.font.underline
                except Exception:
                    pass
        except Exception as e:
            logger.debug(f"Run格式复制失败: {e}")
    
    def _get_images_from_run(self, run, document_part):
        """从run中提取图片数据"""
        images = []
        try:
            import xml.etree.ElementTree as ET
            from io import BytesIO
            
            # 解析run的XML以查找图片
            run_xml = run._element.xml
            root = ET.fromstring(run_xml)
            namespace = {
                'a': "http://schemas.openxmlformats.org/drawingml/2006/main",
                'r': "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
                'wp': "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
            }
            
            # 查找内嵌图片
            inlines = root.findall('.//wp:inline', namespace)
            
            for inline in inlines:
                try:
                    # 获取图片的embed ID
                    blips = inline.findall('.//a:blip', namespace)
                    for blip in blips:
                        embed_id = blip.get(f"{{{namespace['r']}}}embed")
                        if embed_id:
                            try:
                                if hasattr(document_part, 'rels') and embed_id in document_part.rels:
                                    image_part = document_part.rels[embed_id].target_part
                                else:
                                    continue
                            except (TypeError, AttributeError):
                                # 处理不可哈希类型错误
                                continue
                            image_data = image_part.blob
                            
                            # 获取图片尺寸
                            extent = inline.find('.//wp:extent', namespace)
                            width = None
                            height = None
                            if extent is not None:
                                cx = extent.get('cx')
                                cy = extent.get('cy')
                                if cx and cy:
                                    # 转换EMU到Inches (1 inch = 914400 EMU)
                                    from docx.shared import Inches
                                    width = Inches(int(cx) / 914400)
                                    height = Inches(int(cy) / 914400)
                            
                            # 返回给调用者的数据（使用BytesIO）
                            images.append({
                                'stream': BytesIO(image_data),
                                'width': width,
                                'height': height
                            })
                            logger.debug(f"成功提取图片: {embed_id}")
                                
                except Exception as e:
                    logger.warning(f"提取图片数据失败: {e}")
                    continue
                    
        except Exception as e:
            logger.warning(f"解析run中的图片失败: {e}")
            
        return images
    
    def _copy_tables_in_range(self, source_doc: Document, target_doc: Document, paragraph_range: List[int]):
        """复制指定范围内的表格，包括表格中的图片"""
        try:
            # 获取文档中所有元素的顺序（段落和表格）
            document_elements = []
            
            # 遍历文档的body元素，按顺序收集段落和表格
            from docx.oxml.ns import qn
            body = source_doc._body._body
            
            para_index = 0
            table_index = 0
            
            for element in body:
                if element.tag == qn('w:p'):  # 段落
                    document_elements.append(('paragraph', para_index))
                    para_index += 1
                elif element.tag == qn('w:tbl'):  # 表格
                    document_elements.append(('table', table_index))
                    table_index += 1
            
            # 找出在指定段落范围内的表格
            tables_to_copy = []
            for i, (element_type, element_index) in enumerate(document_elements):
                if element_type == 'table':
                    # 检查表格前后的段落是否在范围内
                    should_copy_table = False
                    
                    # 检查表格前面的段落
                    for j in range(i-1, -1, -1):
                        if document_elements[j][0] == 'paragraph':
                            if document_elements[j][1] in paragraph_range:
                                should_copy_table = True
                            break
                    
                    # 如果前面的段落不在范围内，检查表格后面的段落
                    if not should_copy_table:
                        for j in range(i+1, len(document_elements)):
                            if document_elements[j][0] == 'paragraph':
                                if document_elements[j][1] in paragraph_range:
                                    should_copy_table = True
                                break
                    
                    # 如果表格前面或后面的段落在范围内，则复制该表格
                    if should_copy_table:
                        tables_to_copy.append(element_index)
            
            # 复制符合条件的表格
            for table_index in tables_to_copy:
                if table_index < len(source_doc.tables):
                    table = source_doc.tables[table_index]
                    new_table = target_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                    
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            # 复制单元格内容，包括图片
                            self._copy_cell_content(cell, new_table.cell(i, j), source_doc.part)
                    
                    logger.info(f"复制了表格 {table_index}，位于章节范围内")
                        
        except Exception as e:
            logger.warning(f"表格复制警告: {e}")
    
    def _copy_cell_content(self, source_cell, target_cell, source_doc_part):
        """复制表格单元格内容，包括文本和图片"""
        try:
            # 清空目标单元格
            target_cell.text = ""
            
            # 复制每个段落
            for para in source_cell.paragraphs:
                # 如果是第一个段落，使用现有的段落
                if para == source_cell.paragraphs[0] and len(target_cell.paragraphs) > 0:
                    target_para = target_cell.paragraphs[0]
                    # 清空现有内容
                    target_para.clear()
                else:
                    target_para = target_cell.add_paragraph()
                
                # 复制段落格式
                if para.style:
                    try:
                        target_para.style = para.style
                    except Exception:
                        pass
                target_para.alignment = para.alignment
                
                # 复制runs
                for run in para.runs:
                    # 检查run中是否包含图片
                    images_in_run = self._get_images_from_run(run, source_doc_part)
                    
                    if images_in_run:
                        # 如果run包含图片，先添加文本，然后添加图片
                        if run.text:
                            new_run = target_para.add_run(run.text)
                            self._copy_run_format(run, new_run)
                        
                        # 添加图片
                        for image_data in images_in_run:
                            try:
                                new_run = target_para.add_run()
                                new_run.add_picture(image_data['stream'], width=image_data['width'], height=image_data['height'])
                                logger.info(f"成功复制表格中的图片")
                            except Exception as e:
                                logger.warning(f"表格图片复制失败: {e}")
                                # 如果图片复制失败，添加占位文本
                                new_run = target_para.add_run("[图片]")
                    else:
                        # 普通文本run
                        new_run = target_para.add_run(run.text)
                        self._copy_run_format(run, new_run)
                        
        except Exception as e:
            logger.warning(f"单元格内容复制警告: {e}")
    
    def _sanitize_filename(self, filename: str) -> str:
        """清理文件名，移除或替换无效字符"""
        import re
        
        # 移除或替换各种空白字符和特殊字符
        filename = filename.replace('\u3000', ' ')  # 全角空格
        filename = filename.replace('\t', ' ')      # 制表符
        filename = filename.replace('\n', ' ')      # 换行符
        filename = filename.replace('\r', ' ')      # 回车符
        
        # 移除Windows文件名中的无效字符
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        
        # 移除控制字符
        filename = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', filename)
        
        # 压缩多个空格为单个空格
        filename = re.sub(r'\s+', ' ', filename)
        
        # 移除首尾空格
        filename = filename.strip()
        
        # 如果文件名为空，使用默认名称
        if not filename:
            filename = '未命名章节'
        
        # 限制文件名长度（Windows路径限制）
        return filename[:100]

class DocumentProcessor:
    """文档处理器（性能优化版）"""
    
    def __init__(self, input_dir: str, output_dir: str, 
                 file_thread_count: int = 4, chapter_thread_count: int = 2, min_level: int = 5):
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        
        # 根据CPU核心数优化线程数
        import multiprocessing
        cpu_count = multiprocessing.cpu_count()
        self.file_thread_count = min(file_thread_count, cpu_count)
        self.chapter_thread_count = min(chapter_thread_count, max(1, cpu_count // 2))
        
        self.splitter = WordDocumentSplitter(
            min_level=min_level,
            max_workers_docs=self.file_thread_count,
            max_workers_chapters=self.chapter_thread_count
        )
        
        # 创建输出目录
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"处理器配置: 文档线程={self.file_thread_count}, 章节线程={self.chapter_thread_count}, CPU核心={cpu_count}")
    
    def process_all_documents(self):
        """处理所有文档"""
        # 获取所有Word文档
        all_files = list(self.input_dir.glob('*.docx')) + list(self.input_dir.glob('*.doc'))
        
        # 过滤掉以~开头的临时文件（Word打开文档时会创建这些临时文件）
        doc_files = [f for f in all_files if not f.name.startswith('~')]
        
        if not doc_files:
            logger.warning(f"在 {self.input_dir} 中未找到Word文档")
            return
        
        logger.info(f"找到 {len(doc_files)} 个文档，开始处理...")
        
        # 使用线程池处理多个文档
        with ThreadPoolExecutor(max_workers=self.file_thread_count) as executor:
            futures = [executor.submit(self.process_single_document, doc_file) 
                      for doc_file in doc_files]
            
            for future in as_completed(futures):
                try:
                    result = future.result()
                    logger.info(f"文档处理完成: {result}")
                except Exception as e:
                    logger.error(f"文档处理失败: {e}")
    
    def process_single_document(self, doc_path: Path) -> str:
        """处理单个文档（性能优化版）"""
        doc = None
        try:
            logger.info(f"开始处理文档: {doc_path}")
            
            # 分析文档结构
            doc, chapters = self.splitter.analyze_document_structure(str(doc_path))
            
            if not chapters:
                logger.warning(f"文档 {doc_path} 未找到可拆分的章节")
                return f"跳过: {doc_path.name}"
            
            # 创建文档专用的输出目录
            doc_output_dir = self.output_dir / doc_path.stem
            doc_output_dir.mkdir(parents=True, exist_ok=True)
            
            # 动态调整线程数（基于章节数量）
            optimal_workers = min(self.chapter_thread_count, len(chapters), 8)
            
            # 使用线程池处理章节
            with ThreadPoolExecutor(max_workers=optimal_workers) as executor:
                futures = [executor.submit(self.splitter.create_chapter_document, 
                                         doc, chapter, str(doc_output_dir))
                          for chapter in chapters]
                
                created_files = []
                failed_count = 0
                
                for future in as_completed(futures):
                    try:
                        result = future.result(timeout=300)  # 5分钟超时
                        created_files.append(result)
                    except Exception as e:
                        failed_count += 1
                        logger.error(f"章节处理失败: {e}")
                

                
                # 强制垃圾回收
                gc.collect()
            
            result_msg = f"完成: {doc_path.name} -> {len(created_files)} 个章节"
            if failed_count > 0:
                result_msg += f" (失败: {failed_count})"
            
            return result_msg
            
        except Exception as e:
            logger.error(f"处理文档 {doc_path} 时发生错误: {e}")
            return f"失败: {doc_path.name}"
        finally:
            # 确保释放文档对象
            if doc:
                del doc
            gc.collect()

# 模块可以直接导入使用，不包含主函数
# 使用方式：from word_splitter import DocumentProcessor