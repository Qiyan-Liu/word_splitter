#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word Document Splitter Tool
Split large Word documents into multiple smaller documents based on document structure
Supports multi-threading processing while preserving original formatting and styles

author: QiyanLiu
date: 2025-08-06
"""
import os
import sys
import threading
import re
import copy
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import List, Dict, Tuple
import logging
from dataclasses import dataclass
import gc

try:
    from docx import Document
    from docx.shared import Inches
    from docx.oxml.ns import qn
except ImportError:
    print("Please install python-docx library: pip install python-docx")
    sys.exit(1)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('word_splitter.log', encoding='utf-8')
    ]
)
logger = logging.getLogger(__name__)

@dataclass
class ChapterInfo:
    """Chapter information"""
    title: str
    level: int
    start_paragraph: int
    end_paragraph: int
    paragraphs: List[int]

class WordDocumentSplitter:
    """
    Word Document Splitter
    
    Responsible for analyzing Word document structure, identifying chapters, and splitting documents into smaller files.
    Supports multi-threading processing while preserving original formatting and styles.
    """
    
    def __init__(self, min_level: int = 3, max_workers_docs: int = 4, max_workers_chapters: int = 2):
        """
        Initialize document splitter
        
        Args:
            min_level: Minimum split level
            max_workers_docs: Number of threads for processing multiple documents
            max_workers_chapters: Number of threads for processing chapters in a single document
        """
        self.min_level = min_level
        self.max_workers_docs = max_workers_docs
        self.max_workers_chapters = max_workers_chapters
        self.lock = threading.Lock()

    def analyze_document_structure(self, doc_path: str) -> Tuple[Document, List[ChapterInfo]]:
        """Analyze document structure and identify chapters (performance optimized version)
        
        Args:
            doc_path: Word document path
            
        Returns:
            Tuple[Document, List[ChapterInfo]]: Document object and list of chapter information
            
        Raises:
            Exception: Raised when document reading or analysis fails
        """
        import time
        start_time = time.time()
        
        try:
            logger.info(f"Starting document structure analysis: {doc_path}")
            doc = Document(doc_path)
            chapters = []
            current_levels = {}
            all_headings = []
            paragraphs = doc.paragraphs
            total_paragraphs = len(paragraphs)
            logger.info(f"Document contains {total_paragraphs} paragraphs")
            
            for i, paragraph in enumerate(paragraphs):
                outline_level = self._get_outline_level(paragraph)
                text = paragraph.text.strip()
                if outline_level > 0 and text:
                    all_headings.append({
                        'level': outline_level,
                        'title': text,
                        'paragraph_index': i
                    })
            logger.info(f"Found {len(all_headings)} headings")
            
            if all_headings:
                for i, heading in enumerate(all_headings):
                    outline_level = heading['level']
                    paragraph_index = heading['paragraph_index']
                    current_levels[outline_level] = {
                        'title': heading['title'],
                        'paragraph_index': paragraph_index
                    }
                    keys_to_remove = [k for k in current_levels.keys() if k > outline_level]
                    for k in keys_to_remove:
                        del current_levels[k]
                    should_create = self._should_create_chapter_at_position(all_headings, i, outline_level)
                    if should_create:
                        chapter_title = self._build_chapter_title(current_levels, outline_level)
                        chapter = ChapterInfo(
                            title=chapter_title,
                            level=outline_level,
                            start_paragraph=paragraph_index,
                            end_paragraph=paragraph_index,
                            paragraphs=[paragraph_index]
                        )
                        chapters.append(chapter)
            
            self._set_chapter_boundaries(chapters, total_paragraphs, doc)
            end_time = time.time()
            processing_time = end_time - start_time
            logger.info(f"Document structure analysis completed, identified {len(chapters)} chapters, took {processing_time:.2f} seconds")
            return doc, chapters
        except Exception as e:
            logger.error(f"Failed to analyze document structure: {e}")
            raise
    
    def _get_outline_level(self, paragraph) -> int:
        """Get paragraph outline level"""
        return self._calculate_outline_level(paragraph)
    
    def _calculate_outline_level(self, paragraph) -> int:
        """Calculate paragraph outline level (based only on Word standard styles and outline levels)"""
        try:
            # 1. First check if paragraph style is standard heading style
            style_name = paragraph.style.name
            
            # Check English heading styles (Heading 1, Heading 2, etc.)
            if style_name.startswith('Heading '):
                level_str = style_name.replace('Heading ', '')
                try:
                    return int(level_str)
                except ValueError:
                    return 1  # Default to level 1 heading
            
            # Check Chinese heading styles (标题 1, 标题 2, etc.)
            if '标题' in style_name:
                level_match = re.search(r'(\d+)', style_name)
                if level_match:
                    return int(level_match.group(1))
                return 1  # Default to level 1 heading
            
            # Check custom styles (样式1, 样式2, 样式3, etc.)
            if style_name.startswith('样式'):
                level_match = re.search(r'(\d+)', style_name)
                if level_match:
                    level_num = int(level_match.group(1))
                    # 样式3 is usually used as level 1 heading, 样式4 as level 2 heading, etc.
                    # But to be safe, we check paragraph content to determine
                    text = paragraph.text.strip()
                    if text and self._looks_like_heading(text):
                        return level_num
                return 0  # Not a heading style
            
            # 2. Check paragraph format outline level (Word built-in outline level attribute)
            if hasattr(paragraph._element, 'pPr') and paragraph._element.pPr is not None:
                outline_lvl = paragraph._element.pPr.find(qn('w:outlineLvl'))
                if outline_lvl is not None:
                    level_value = int(outline_lvl.get(qn('w:val')))
                    return level_value + 1  # Word outline levels start from 0, we start from 1
            
            # 3. If neither heading style nor outline level, then not a heading
            return 0
            
        except Exception:
            return 0
    
    def _looks_like_heading(self, text: str) -> bool:
        """Determine if text looks like a heading"""
        if not text or len(text.strip()) == 0:
            return False
        
        text = text.strip()
        
        # Check if it's chapter heading format (一、二、三、etc.)
        if re.match(r'^[一二三四五六七八九十]+、', text):
            return True
        
        # Check if it's numeric chapter format (1、2、3、etc.)
        if re.match(r'^\d+[、.]', text):
            return True
        
        # Check length (headings are usually short)
        if len(text) <= 50 and not '。' in text:
            # Short text without period, might be a heading
            return True
        
        return False
    
    def _is_likely_toc_content(self, text: str) -> bool:
        """Determine if it's table of contents content"""
        # TOC characteristics: contains page numbers, dot connections, specific keywords, etc.
        toc_indicators = [
            '目　　录',
            '目录',
            '......',
            '………',
            '.....',
            '-----',
        ]
        
        # Check if contains TOC characteristics
        for indicator in toc_indicators:
            if indicator in text:
                return True
        
        # Check if it's page number format (ending with numbers)
        if re.search(r'\d+\s*$', text):
            return True
        
        # Check if contains many spaces (TOC alignment format)
        if text.count('　') > 2 or text.count(' ') > 10:
            return True
            
        return False
    
    def _is_likely_heading(self, paragraph, text: str) -> bool:
        """Determine if it's a real heading"""
        # Heading characteristics check
        if not paragraph.runs:
            return False
        
        # Exclude content that's obviously not headings
        
        # Exclude table-related content (descriptive text ending with colon)
        if re.search(r'(表格|图片|图表).*[：:]$', text):
            return False
        
        # Exclude ending statements
        if re.search(r'(结束|完毕|完成)[。.]$', text):
            return False
        
        # Exclude overly long text (headings are usually short)
        if len(text) > 50:
            return False
        
        # Check font bold
        is_bold = any(run.bold for run in paragraph.runs if run.bold is not None)
        
        # Check font size
        font_size = paragraph.runs[0].font.size
        is_large_font = font_size and font_size.pt >= 12
        
        # Check if contains chapter numbering
        has_chapter_number = bool(re.search(r'^(第[一二三四五六七八九十\d]+[章节部分]|\d+[\..]\d*\s*|[一二三四五六七八九十]+[、\.]\s*)', text))
        
        # Comprehensive judgment
        score = 0
        if is_bold:
            score += 2
        if is_large_font:
            score += 1
        if has_chapter_number:
            score += 3  # Chapter numbering has higher weight
        
        # Lower threshold, score>=3 considered as heading (balance strictness and recognition rate)
        return score >= 3
    
    # ========================================================================
    # Chapter creation and boundary setting methods
    # ========================================================================
    
    def _should_create_chapter_at_position(self, all_headings: List[Dict], current_index: int, current_level: int) -> bool:
        """Determine whether to create chapter at current position (locally adaptive chapter levels)"""
        # Analyze maximum level of current chapter branch
        branch_max_level = self._get_branch_max_level(all_headings, current_index, current_level)
        
        # If current branch max level is less than min_level, use branch max level as split level
        effective_min_level = min(self.min_level, branch_max_level)
        
        # Only create chapter when reaching effective minimum level
        if current_level == effective_min_level:
            return True
        
        return False
    
    def _get_branch_max_level(self, all_headings: List[Dict], current_index: int, current_level: int) -> int:
        """Get maximum level of current chapter branch"""
        # Build current chapter path (from root to current position)
        current_path = {}
        for i in range(current_index + 1):
            heading = all_headings[i]
            if heading['level'] <= current_level:
                current_path[heading['level']] = heading['title']
            if heading['level'] == current_level and i == current_index:
                break
        
        # Find all subsequent headings belonging to current branch and find maximum level
        max_level = current_level
        
        for i in range(current_index + 1, len(all_headings)):
            heading = all_headings[i]
            
            # If encountering same level or higher level heading, current branch ends
            if heading['level'] <= current_level:
                break
            
            # Check if belongs to current branch
            if self._belongs_to_current_branch(all_headings, i, current_path, current_level):
                max_level = max(max_level, heading['level'])
        
        return max_level
    
    def _belongs_to_current_branch(self, all_headings: List[Dict], heading_index: int, current_path: Dict, current_level: int) -> bool:
        """判断指定的标题是否属于当前分支"""
        # Rebuild path to specified heading
        path_to_heading = {}
        
        for i in range(heading_index + 1):
            heading = all_headings[i]
            if heading['level'] <= current_level:
                path_to_heading[heading['level']] = heading['title']
        
        # Check if paths match
        for level in range(1, current_level + 1):
            if level in current_path and level in path_to_heading:
                if current_path[level] != path_to_heading[level]:
                    return False
            elif level in current_path or level in path_to_heading:
                return False
        
        return True
    

    
    def _has_deeper_branch(self, all_headings: List[Dict], current_index: int, current_level: int) -> bool:
        """检查当前分支是否还有更深的层级"""
        # Build current path
        current_path = {}
        for i in range(current_index + 1):
            heading = all_headings[i]
            if heading['level'] <= current_level:
                current_path[heading['level']] = heading['title']
            if heading['level'] == current_level and i == current_index:
                break
        
        # Check if there are deeper levels belonging to same branch afterwards
        for i in range(current_index + 1, len(all_headings)):
            future_heading = all_headings[i]
            future_level = future_heading['level']
            
            # If encountering same level or shallower level, current branch ends
            if future_level <= current_level:
                break
            
            # If it's deeper level, check if belongs to current branch
            if future_level > current_level:
                # Build path to this deeper level
                future_path = {}
                for j in range(current_index + 1):
                    heading = all_headings[j]
                    if heading['level'] <= current_level:
                        future_path[heading['level']] = heading['title']
                
                # Check if paths match (i.e., belongs to same branch)
                path_matches = True
                for level in current_path:
                    if level not in future_path or current_path[level] != future_path[level]:
                        path_matches = False
                        break
                
                if path_matches:
                    return True
        
        return False
    
    def _determine_target_level(self, unique_levels: List[int]) -> int:
        """根据文档的整体结构确定目标拆分层级（保留用于向后兼容）"""
        if not unique_levels:
            return self.min_level
        
        # Check if levels are continuous
        is_continuous = True
        if len(unique_levels) > 1:
            for i in range(len(unique_levels) - 1):
                if unique_levels[i + 1] - unique_levels[i] > 1:
                    is_continuous = False
                    break
        
        max_level = max(unique_levels)
        
        if not is_continuous:
            # Levels not continuous, find deepest level less than min_level
            valid_levels = [level for level in unique_levels if level < self.min_level]
            if valid_levels:
                return max(valid_levels)
            else:
                # If no level less than min_level, return minimum level
                return min(unique_levels)
        
        # Case when levels are continuous
        # If max level reaches or exceeds min_level, create chapter at min_level
        if max_level >= self.min_level:
            return self.min_level
        
        # If max level is less than min_level, create chapter at deepest level (but at least level 2)
        return max_level if max_level >= 2 else 2
    
    def _should_create_chapter(self, current_levels: Dict, outline_level: int) -> bool:
        """判断是否应该创建新章节"""
        # Get currently existing levels
        existing_levels = sorted(current_levels.keys())
        
        if not existing_levels:
            return False
        
        # Check if levels are continuous
        is_continuous = True
        if len(existing_levels) > 1:
            for i in range(len(existing_levels) - 1):
                if existing_levels[i + 1] - existing_levels[i] > 1:
                    is_continuous = False
                    break
        
        max_level = max(existing_levels)
        
        if not is_continuous:
            # Levels not continuous, find deepest level less than min_level
            valid_levels = [level for level in existing_levels if level < self.min_level]
            if valid_levels:
                target_level = max(valid_levels)
                return outline_level == target_level and target_level >= 2
            else:
                # If no level less than min_level, don't create chapter
                return False
        
        # Case when levels are continuous
        # If max level reaches or exceeds min_level, only create chapter at min_level
        if max_level >= self.min_level:
            return outline_level == self.min_level
        
        # If max level is less than min_level, create chapter at deepest level
        return outline_level == max_level and max_level >= 2
    
    def _build_chapter_title(self, current_levels: Dict, target_level: int) -> str:
        """构建章节标题"""
        titles = []
        for level in sorted(current_levels.keys()):
            if level <= target_level:
                titles.append(current_levels[level]['title'])
        return ' - '.join(titles) if titles else f"章节_{target_level}"
    
    def _set_chapter_boundaries(self, chapters: List[ChapterInfo], total_paragraphs: int, doc: Document = None):
        """设置章节边界（智能处理层级跳跃）"""
        if not chapters:
            return
        
        # Get position information of all headings in document
        all_heading_positions = []
        
        if doc:
            for i, paragraph in enumerate(doc.paragraphs):
                outline_level = self._get_outline_level(paragraph)
                text = paragraph.text.strip()
                if outline_level > 0 and text:
                    all_heading_positions.append({
                        'paragraph_index': i,
                        'level': outline_level,
                        'title': text
                    })
        
        for i, chapter in enumerate(chapters):
            if i < len(chapters) - 1:
                next_chapter = chapters[i + 1]
                
                # Intelligently determine chapter end position
                # Find if there are same level or higher level headings between current and next chapter
                chapter_end = next_chapter.start_paragraph - 1
                
                # Check for content inclusion issues caused by level jumps
                for heading_info in all_heading_positions:
                    heading_pos = heading_info['paragraph_index']
                    heading_level = heading_info['level']
                    
                    # If same level or higher level heading found within current chapter range
                    if (chapter.start_paragraph < heading_pos < next_chapter.start_paragraph and 
                        heading_level <= chapter.level):
                        # Adjust chapter end position to before this heading
                        chapter_end = heading_pos - 1
                        break
                
                chapter.end_paragraph = max(chapter.start_paragraph, chapter_end)
            else:
                # Last chapter
                chapter.end_paragraph = total_paragraphs - 1
            
            # Set paragraphs included in chapter
            chapter.paragraphs = list(range(chapter.start_paragraph, chapter.end_paragraph + 1))
    
    # ========================================================================
    # 文档创建和内容复制方法
    # ========================================================================
    
    def create_chapter_document(self, original_doc: Document, chapter: ChapterInfo, output_path: str):
        """
        Create chapter document (performance optimized version)
        
        Args:
            original_doc: Original document object
            chapter: Chapter information
            output_path: Output directory path
            
        Returns:
            str: Full path of created document
            
        Raises:
            Exception: Raised when document creation or saving fails
        """
        try:
            # Create new document
            new_doc = Document()
            
            # Copy document-level settings (including theme fonts)
            self._copy_document_settings(original_doc, new_doc)
            
            # Copy styles
            self._copy_styles(original_doc, new_doc)
            
            # Batch copy chapter content (reduce single call overhead)
            self._copy_paragraphs_batch(original_doc, new_doc, chapter.paragraphs)
            
            # Copy tables (if within chapter range)
            self._copy_tables_in_range(original_doc, new_doc, chapter.paragraphs)
            
            # Save document
            safe_filename = self._sanitize_filename(chapter.title)
            full_output_path = os.path.join(output_path, f"{safe_filename}.docx")
            
            with self.lock:
                os.makedirs(os.path.dirname(full_output_path), exist_ok=True)
                new_doc.save(full_output_path)
            
            # Force garbage collection to free memory
            del new_doc
            gc.collect()
            
            logger.info(f"Chapter document created: {full_output_path}")
            return full_output_path
            
        except Exception as e:
            logger.error(f"Failed to create chapter document: {e}")
            raise
    
    def _copy_paragraphs_batch(self, source_doc: Document, target_doc: Document, paragraph_indices: List[int]):
        """Batch copy paragraphs (performance optimized)"""
        try:
            source_paragraphs = source_doc.paragraphs
            source_part = source_doc.part
            
            # Batch process paragraphs, reduce function call overhead
            for para_index in paragraph_indices:
                if para_index < len(source_paragraphs):
                    original_para = source_paragraphs[para_index]
                    self._copy_paragraph(original_para, target_doc, source_part)
                    
        except Exception as e:
            logger.warning(f"Batch copy paragraphs failed: {e}")
            # Fallback to individual copying
            for para_index in paragraph_indices:
                if para_index < len(source_doc.paragraphs):
                    try:
                        original_para = source_doc.paragraphs[para_index]
                        self._copy_paragraph(original_para, target_doc, source_doc.part)
                    except Exception as inner_e:
                        logger.warning(f"Failed to copy paragraph {para_index}: {inner_e}")
                        continue
    
    def _copy_document_settings(self, source_doc: Document, target_doc: Document):
        """Copy document-level settings to ensure correct font display"""
        try:
            # Method 1: Set default font for Normal style
            try:
                # Extract main font from source document
                main_font = self._extract_main_font(source_doc)
                if main_font:
                    normal_style = target_doc.styles['Normal']
                    normal_style.font.name = main_font
                    logger.info(f"Normal style font set: {main_font}")
            except Exception as e:
                logger.warning(f"Failed to set Normal style font: {e}")
            
            # Method 2: Set document-level default font
            try:
                main_font = self._extract_main_font(source_doc)
                if main_font:
                    # Get document XML
                    doc_xml = target_doc.part.blob.decode('utf-8')
                    
                    # Create document default font settings
                    doc_defaults = f'''
<w:docDefaults>
    <w:rPrDefault>
        <w:rPr>
            <w:rFonts w:ascii="{main_font}" w:eastAsia="{main_font}" w:hAnsi="{main_font}" w:cs="{main_font}"/>
        </w:rPr>
    </w:rPrDefault>
    <w:pPrDefault/>
</w:docDefaults>'''
                    
                    # Insert docDefaults before styles
                    if '<w:docDefaults>' not in doc_xml:
                        styles_pattern = r'<w:styles[^>]*>'
                        doc_xml = re.sub(styles_pattern, doc_defaults + '\n\g<0>', doc_xml)
                        
                        # Update document
                        target_doc.part._blob = doc_xml.encode('utf-8')
                        logger.info(f"Document default font set: {main_font}")
                        
            except Exception as e:
                logger.warning(f"Failed to set document default font: {e}")
                
        except Exception as e:
            logger.warning(f"Failed to copy document settings: {e}")
    
    def _extract_main_font(self, doc: Document) -> str:
        """Extract main font from document"""
        try:
            # Traverse document paragraphs, find first run with font settings
            for para in doc.paragraphs:
                if para.text.strip() and not para.style.name.startswith('Heading'):
                    for run in para.runs:
                        if run.font.name:
                            return run.font.name
                        # Check rFonts attribute
                        if hasattr(run.font._element, 'rFonts'):
                            rfonts = run.font._element.rFonts
                            if rfonts is not None:
                                if rfonts.eastAsia:
                                    return rfonts.eastAsia
                                if rfonts.ascii:
                                    return rfonts.ascii
                                if rfonts.hAnsi:
                                    return rfonts.hAnsi
            
            # If not found, return default font
            return "微软雅黑"
            
        except Exception as e:
            logger.warning(f"Failed to extract main font: {e}")
            return "微软雅黑"
    
    # ========================================================================
    # Style and format copying methods
    # ========================================================================
    
    def _copy_styles(self, source_doc: Document, target_doc: Document):
        """Copy document styles"""
        try:
            # Get existing style names in target document (one-time fetch to avoid repeated calculations)
            existing_styles = {s.name for s in target_doc.styles}
            
            # Copy paragraph styles
            for style in source_doc.styles:
                if style.name not in existing_styles:
                    try:
                        new_style = target_doc.styles.add_style(style.name, style.type)
                        
                        # Copy style font attributes - improved for Chinese fonts
                        if hasattr(style, 'font') and hasattr(new_style, 'font'):
                            try:
                                # Copy basic font name
                                if style.font.name:
                                    new_style.font.name = style.font.name
                                
                                # Copy detailed font information from XML for better Chinese font support
                                if hasattr(style.font, '_element') and hasattr(new_style.font, '_element'):
                                    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                                    source_rfonts = style.font._element.find(f'.//{{{w_ns}}}rFonts')
                                    if source_rfonts is not None:
                                        # Get or create target rFonts element
                                        target_rfonts = new_style.font._element.find(f'.//{{{w_ns}}}rFonts')
                                        if target_rfonts is None:
                                            from xml.etree.ElementTree import Element
                                            target_rfonts = Element(f'{{{w_ns}}}rFonts')
                                            new_style.font._element.append(target_rfonts)
                                        
                                        # Copy all font attributes including theme fonts
                                        font_attrs = ['ascii', 'eastAsia', 'hAnsi', 'cs', 'asciiTheme', 'eastAsiaTheme', 'hAnsiTheme', 'csTheme']
                                        for attr in font_attrs:
                                            font_value = source_rfonts.get(f'{{{w_ns}}}{attr}')
                                            if font_value:
                                                target_rfonts.set(f'{{{w_ns}}}{attr}', font_value)
                                
                                if style.font.size:
                                    new_style.font.size = style.font.size
                                if style.font.bold is not None:
                                    new_style.font.bold = style.font.bold
                                if style.font.italic is not None:
                                    new_style.font.italic = style.font.italic
                                if style.font.underline is not None:
                                    new_style.font.underline = style.font.underline
                                if style.font.color and style.font.color.rgb:
                                    new_style.font.color.rgb = style.font.color.rgb
                            except Exception as e:
                                logger.debug(f"Failed to copy style font attributes: {e}")
                        
                        # Copy paragraph format attributes
                        if hasattr(style, 'paragraph_format') and hasattr(new_style, 'paragraph_format'):
                            try:
                                src_pf = style.paragraph_format
                                new_pf = new_style.paragraph_format
                                
                                if src_pf.alignment is not None:
                                    new_pf.alignment = src_pf.alignment
                                if src_pf.left_indent:
                                    new_pf.left_indent = src_pf.left_indent
                                if src_pf.right_indent:
                                    new_pf.right_indent = src_pf.right_indent
                                if src_pf.first_line_indent:
                                    new_pf.first_line_indent = src_pf.first_line_indent
                                if src_pf.space_before:
                                    new_pf.space_before = src_pf.space_before
                                if src_pf.space_after:
                                    new_pf.space_after = src_pf.space_after
                                if src_pf.line_spacing:
                                    new_pf.line_spacing = src_pf.line_spacing
                                if src_pf.line_spacing_rule:
                                    new_pf.line_spacing_rule = src_pf.line_spacing_rule
                            except Exception as e:
                                logger.debug(f"Failed to copy style paragraph format: {e}")
                                
                    except Exception as e:
                        logger.debug(f"Failed to copy style {style.name}: {e}")
                else:
                    # If style already exists, update its attributes to ensure consistency
                    try:
                        existing_style = None
                        for s in target_doc.styles:
                            if s.name == style.name:
                                existing_style = s
                                break
                        
                        if existing_style and hasattr(style, 'font') and hasattr(existing_style, 'font'):
                            # Update existing style font attributes
                            if style.font.name and not existing_style.font.name:
                                existing_style.font.name = style.font.name
                            if style.font.size and not existing_style.font.size:
                                existing_style.font.size = style.font.size
                    except Exception as e:
                        logger.debug(f"Failed to update existing style: {e}")
            
            # Ensure Hyperlink style exists in target document
            self._ensure_hyperlink_style(target_doc)
            
        except Exception as e:
            logger.warning(f"Style copy warning: {e}")
    
    def _copy_paragraph(self, source_para, target_doc: Document, source_doc_part):
        """Copy paragraph including text, format and images"""
        try:
            new_para = target_doc.add_paragraph()
            
            # Safely copy paragraph format
            try:
                if hasattr(source_para, 'style') and source_para.style:
                    style_name = str(source_para.style.name) if source_para.style.name else None
                    if style_name:
                        # Check if target document has this style
                        style_found = False
                        for style in target_doc.styles:
                            if style.name == style_name:
                                new_para.style = style
                                style_found = True
                                break
                        
                        # If style doesn't exist, try to copy direct format from original paragraph
                        if not style_found:
                            logger.debug(f"Style {style_name} not found, will copy direct format")
                            # Try to get font information from source paragraph style
                            try:
                                if hasattr(source_para.style, 'font') and source_para.style.font:
                                    source_font = source_para.style.font
                                    # Set default font for all runs in paragraph
                                    if hasattr(source_font, 'name') and source_font.name:
                                        # Record style default font, apply to run later
                                        setattr(new_para, '_default_font_name', source_font.name)
                                    if hasattr(source_font, 'size') and source_font.size:
                                        setattr(new_para, '_default_font_size', source_font.size)
                                else:
                                    # If style has no font info, try to get font from first run of source paragraph
                                    if source_para.runs:
                                        first_run = source_para.runs[0]
                                        if (hasattr(first_run, 'font') and first_run.font and 
                                            hasattr(first_run.font, 'name') and first_run.font.name):
                                            setattr(new_para, '_default_font_name', first_run.font.name)
                                        if (hasattr(first_run, 'font') and first_run.font and 
                                            hasattr(first_run.font, 'size') and first_run.font.size):
                                            setattr(new_para, '_default_font_size', first_run.font.size)
                            except Exception as e:
                                logger.debug(f"Failed to copy style font: {e}")
                        else:
                            # Even if style exists, check if default font needs to be set
                            try:
                                if (hasattr(source_para.style, 'font') and source_para.style.font and 
                                    hasattr(source_para.style.font, 'name') and source_para.style.font.name):
                                    setattr(new_para, '_default_font_name', source_para.style.font.name)
                                elif source_para.runs:
                                    # Get font from first run as default font
                                    first_run = source_para.runs[0]
                                    if (hasattr(first_run, 'font') and first_run.font and 
                                        hasattr(first_run.font, 'name') and first_run.font.name):
                                        setattr(new_para, '_default_font_name', first_run.font.name)
                            except Exception as e:
                                logger.debug(f"Failed to set default font: {e}")
            except Exception as e:
                logger.debug(f"Failed to copy style: {e}")
            
            # Safely copy paragraph format
            try:
                if hasattr(source_para, 'alignment') and source_para.alignment is not None:
                    new_para.alignment = source_para.alignment
            except Exception as e:
                logger.debug(f"Failed to copy alignment: {e}")
            
            # Copy paragraph indentation
            try:
                if hasattr(source_para, 'paragraph_format') and source_para.paragraph_format:
                    pf = source_para.paragraph_format
                    new_pf = new_para.paragraph_format
                    
                    # Left indent
                    if hasattr(pf, 'left_indent') and pf.left_indent:
                        new_pf.left_indent = pf.left_indent
                    
                    # Right indent
                    if hasattr(pf, 'right_indent') and pf.right_indent:
                        new_pf.right_indent = pf.right_indent
                    
                    # First line indent
                    if hasattr(pf, 'first_line_indent') and pf.first_line_indent:
                        new_pf.first_line_indent = pf.first_line_indent
                    
                    # Space before paragraph
                    if hasattr(pf, 'space_before') and pf.space_before:
                        new_pf.space_before = pf.space_before
                    
                    # Space after paragraph
                    if hasattr(pf, 'space_after') and pf.space_after:
                        new_pf.space_after = pf.space_after
                    
                    # Line spacing
                    if hasattr(pf, 'line_spacing') and pf.line_spacing:
                        new_pf.line_spacing = pf.line_spacing
                    
                    # Line spacing rule
                    if hasattr(pf, 'line_spacing_rule') and pf.line_spacing_rule:
                        new_pf.line_spacing_rule = pf.line_spacing_rule
                    
                    # Widow control
                    if hasattr(pf, 'widow_control') and pf.widow_control is not None:
                        new_pf.widow_control = pf.widow_control
                    
                    # Keep with next paragraph
                    if hasattr(pf, 'keep_with_next') and pf.keep_with_next is not None:
                        new_pf.keep_with_next = pf.keep_with_next
                    
                    # Keep together
                    if hasattr(pf, 'keep_together') and pf.keep_together is not None:
                        new_pf.keep_together = pf.keep_together
                    
                    # Page break before
                    if hasattr(pf, 'page_break_before') and pf.page_break_before is not None:
                        new_pf.page_break_before = pf.page_break_before
            except Exception as e:
                logger.debug(f"Failed to copy paragraph format: {e}")
            
            # Check if paragraph contains hyperlinks that need special handling
            has_hyperlinks = self._paragraph_has_hyperlinks(source_para)
            
            if has_hyperlinks:
                # Handle hyperlinks separately to avoid duplication
                self._copy_paragraph_hyperlinks(source_para, new_para, source_doc_part)
            else:
                # Copy runs normally for paragraphs without hyperlinks
                for run in source_para.runs:
                    try:
                        # Check if run contains images
                        images_in_run = self._get_images_from_run(run, source_doc_part)
                        
                        if images_in_run:
                            # If run contains images, add text first, then add images
                            if run.text:
                                new_run = new_para.add_run(run.text)
                                self._copy_run_format(run, new_run, source_para, new_para)
                                # Copy hyperlink if exists
                                self._copy_hyperlink(run, new_run, source_para, new_para)
                            
                            # Add images
                            for image_data in images_in_run:
                                try:
                                    new_run = new_para.add_run()
                                    new_run.add_picture(image_data['stream'], width=image_data['width'], height=image_data['height'])
                                except Exception as e:
                                    logger.warning(f"Failed to copy image: {e}")
                                    # If image copy fails, add placeholder text
                                    new_run = new_para.add_run("[Image]")
                        else:
                            # Regular text run
                            if run.text:  # Only copy runs with text
                                new_run = new_para.add_run(run.text)
                                self._copy_run_format(run, new_run, source_para, new_para)
                                # Copy hyperlink if exists
                                self._copy_hyperlink(run, new_run, source_para, new_para)
                            
                                # If run has no font name but paragraph has default font, apply default font
                                if (not new_run.font.name and 
                                    hasattr(new_para, '_default_font_name') and 
                                    new_para._default_font_name):
                                    new_run.font.name = new_para._default_font_name
                    except Exception as e:
                        logger.debug(f"Failed to copy run: {e}")
                        continue
                
        except Exception as e:
            logger.warning(f"Paragraph copy warning: {e}")
    
    def _paragraph_has_hyperlinks(self, source_para):
        """Check if paragraph contains hyperlinks (both standard and HYPERLINK fields)"""
        try:
            if hasattr(source_para, '_element') and source_para._element is not None:
                w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                
                # Check for standard w:hyperlink elements
                hyperlink_elems = source_para._element.findall(f'.//{{{w_ns}}}hyperlink')
                if hyperlink_elems:
                    return True
                
                # Check for HYPERLINK field codes
                instr_text_elems = source_para._element.findall(f'.//{{{w_ns}}}instrText')
                for instr_elem in instr_text_elems:
                    if instr_elem.text and 'HYPERLINK' in instr_elem.text:
                        return True
                        
            return False
        except Exception as e:
            logger.debug(f"Failed to check hyperlinks: {e}")
            return False
    
    def _copy_paragraph_hyperlinks(self, source_para, target_para, source_doc_part):
        """Copy all content from paragraph with hyperlinks, maintaining visual text order"""
        try:
            if hasattr(source_para, '_element') and source_para._element is not None:
                # Get hyperlink information from source paragraph first
                hyperlink_info = self._extract_hyperlink_info(source_para)
                
                # Copy only runs that contain actual text, skip empty runs
                # Store run information for later hyperlink application
                copied_runs = []
                for run in source_para.runs:
                    if run.text.strip():  # Only copy runs with non-empty text
                        new_run = target_para.add_run(run.text)
                        self._copy_run_format(run, new_run, source_para, target_para)
                        copied_runs.append((new_run, run.text.strip()))
                    
                    # Handle images in run
                    images_in_run = self._get_images_from_run(run, source_doc_part)
                    for image_data in images_in_run:
                        try:
                            img_run = target_para.add_run()
                            img_run.add_picture(image_data['stream'], width=image_data['width'], height=image_data['height'])
                        except Exception as e:
                            logger.warning(f"Failed to copy image: {e}")
                            img_run = target_para.add_run("[Image]")
                
                # Apply hyperlinks after all runs are copied to maintain order
                for new_run, run_text in copied_runs:
                    for hyperlink in hyperlink_info:
                        if run_text == hyperlink['text'].strip():
                            self._add_hyperlink_to_run(new_run, hyperlink['url'], run_text)
                            break
                        
        except Exception as e:
            logger.debug(f"Failed to copy paragraph hyperlinks: {e}")
    
    def _extract_hyperlink_info(self, source_para):
        """Extract hyperlink information from source paragraph"""
        hyperlink_info = []
        try:
            if hasattr(source_para, '_element') and source_para._element is not None:
                w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                
                # Find all instrText elements that contain HYPERLINK
                instr_text_elems = source_para._element.findall(f'.//{{{w_ns}}}instrText')
                
                for instr_elem in instr_text_elems:
                    if instr_elem.text and 'HYPERLINK' in instr_elem.text:
                        try:
                            # Extract URL from HYPERLINK field
                            hyperlink_text = instr_elem.text
                            import re
                            url_match = re.search(r'HYPERLINK\s+"([^"]+)"', hyperlink_text)
                            
                            if url_match:
                                url = url_match.group(1)
                                logger.debug(f"Found HYPERLINK field with URL: {url}")
                                
                                # Extract the display text for this hyperlink
                                display_text = self._extract_hyperlink_field_text(source_para._element, instr_elem)
                                
                                if display_text:
                                    hyperlink_info.append({
                                        'text': display_text,
                                        'url': url
                                    })
                                    logger.debug(f"Extracted hyperlink: text='{display_text}', url='{url}'")
                                    
                        except Exception as e:
                            logger.debug(f"Failed to process HYPERLINK field: {e}")
                            
        except Exception as e:
            logger.debug(f"Failed to extract hyperlink info: {e}")
        
        return hyperlink_info
    
    def _extract_hyperlink_field_text(self, para_elem, instr_elem):
        """Extract the display text from a HYPERLINK field"""
        try:
            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            
            # Find the run containing the instrText
            instr_run = instr_elem.getparent()
            if instr_run is None:
                return None
            
            # Find all runs in the paragraph
            all_runs = para_elem.findall(f'.//{{{w_ns}}}r')
            instr_run_index = -1
            
            # Find the index of the run containing instrText
            for i, run in enumerate(all_runs):
                if instr_elem in run:
                    instr_run_index = i
                    break
            
            if instr_run_index == -1:
                return None
            
            # Look for fldChar with fldCharType="separate" after the instrText
            # The text after "separate" and before "end" is the display text
            hyperlink_texts = []
            found_separate = False
            
            for i in range(instr_run_index, len(all_runs)):
                run = all_runs[i]
                
                # Check for fldChar elements
                fld_chars = run.findall(f'{{{w_ns}}}fldChar')
                for fld_char in fld_chars:
                    fld_char_type = fld_char.get(f'{{{w_ns}}}fldCharType')
                    if fld_char_type == 'separate':
                        found_separate = True
                    elif fld_char_type == 'end' and found_separate:
                        # End of field, return collected text
                        return ''.join(hyperlink_texts)
                
                # If we found separate, collect text from subsequent runs
                if found_separate:
                    # Check if this run has fldChar end, if so, don't include text after it
                    has_end = any(fc.get(f'{{{w_ns}}}fldCharType') == 'end' for fc in fld_chars)
                    if has_end:
                        # This run contains the end marker, don't include its text
                        break
                    
                    # Collect text from this run
                    text_elems = run.findall(f'{{{w_ns}}}t')
                    for text_elem in text_elems:
                        if text_elem.text:
                            hyperlink_texts.append(text_elem.text)
            
            return ''.join(hyperlink_texts) if hyperlink_texts else None
            
        except Exception as e:
            logger.debug(f"Failed to extract hyperlink field text: {e}")
            return None
    
    def _copy_hyperlink_field_formatting(self, para_elem, instr_elem, target_run):
        """Copy formatting from hyperlink field text to target run"""
        try:
            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            
            # Find the run containing the instrText
            instr_run = instr_elem.getparent()
            if instr_run is None:
                return
            
            # Find all runs in the paragraph
            all_runs = para_elem.findall(f'.//{{{w_ns}}}r')
            instr_run_index = -1
            
            # Find the index of the run containing instrText
            for i, run in enumerate(all_runs):
                if instr_elem in run:
                    instr_run_index = i
                    break
            
            if instr_run_index == -1:
                return
            
            # Look for the first text run after fldChar separate
            found_separate = False
            
            for i in range(instr_run_index, len(all_runs)):
                run = all_runs[i]
                
                # Check for fldChar elements
                fld_chars = run.findall(f'{{{w_ns}}}fldChar')
                for fld_char in fld_chars:
                    fld_char_type = fld_char.get(f'{{{w_ns}}}fldCharType')
                    if fld_char_type == 'separate':
                        found_separate = True
                    elif fld_char_type == 'end' and found_separate:
                        return
                
                # If we found separate and this run has text, copy its formatting
                if found_separate:
                    text_elems = run.findall(f'{{{w_ns}}}t')
                    if text_elems:
                        # Copy formatting from this run
                        rpr_elem = run.find(f'{{{w_ns}}}rPr')
                        if rpr_elem is not None:
                            self._apply_run_format_from_xml(target_run, rpr_elem)
                        return
            
        except Exception as e:
            logger.debug(f"Failed to copy hyperlink field formatting: {e}")
    
    def _apply_run_format_from_xml(self, run, rpr_elem):
        """Apply run formatting from XML rPr element"""
        try:
            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            
            # Font name
            rfonts_elem = rpr_elem.find(f'{{{w_ns}}}rFonts')
            if rfonts_elem is not None:
                ascii_font = rfonts_elem.get(f'{{{w_ns}}}ascii')
                if ascii_font:
                    run.font.name = ascii_font
            
            # Font size
            sz_elem = rpr_elem.find(f'{{{w_ns}}}sz')
            if sz_elem is not None:
                sz_val = sz_elem.get(f'{{{w_ns}}}val')
                if sz_val:
                    try:
                        # Convert half-points to points
                        from docx.shared import Pt
                        run.font.size = Pt(int(sz_val) / 2)
                    except ValueError:
                        pass
            
            # Bold
            b_elem = rpr_elem.find(f'{{{w_ns}}}b')
            if b_elem is not None:
                run.font.bold = True
            
            # Italic
            i_elem = rpr_elem.find(f'{{{w_ns}}}i')
            if i_elem is not None:
                run.font.italic = True
            
            # Underline
            u_elem = rpr_elem.find(f'{{{w_ns}}}u')
            if u_elem is not None:
                run.font.underline = True
                
        except Exception as e:
            logger.debug(f"Failed to apply run format from XML: {e}")
    
    def _copy_run_format(self, source_run, target_run, source_para=None, target_para=None):
        """Safely copy run format"""
        try:
            # Safely copy font format
            if hasattr(source_run, 'font') and source_run.font:
                # Font name - complete copy of rFonts attributes and direct font name
                try:
                    # First try to copy direct font name
                    if hasattr(source_run.font, 'name') and source_run.font.name:
                        target_run.font.name = str(source_run.font.name)
                    
                    # Then copy rFonts attributes to ensure completeness
                    if hasattr(source_run.font, '_element') and hasattr(target_run.font, '_element'):
                        w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                        source_rfonts = source_run.font._element.find(f'.//{{{w_ns}}}rFonts')
                        if source_rfonts is not None:
                            # Get or create target rFonts element
                            target_rfonts = target_run.font._element.find(f'.//{{{w_ns}}}rFonts')
                            if target_rfonts is None:
                                # Create new rFonts element
                                from xml.etree.ElementTree import Element
                                target_rfonts = Element(f'{{{w_ns}}}rFonts')
                                target_run.font._element.append(target_rfonts)
                            
                            # Copy font attributes one by one, ensure eastAsia and theme fonts are not lost
                            font_attrs = ['ascii', 'eastAsia', 'hAnsi', 'cs', 'asciiTheme', 'eastAsiaTheme', 'hAnsiTheme', 'csTheme']
                            for attr in font_attrs:
                                font_value = source_rfonts.get(f'{{{w_ns}}}{attr}')
                                if font_value:
                                    target_rfonts.set(f'{{{w_ns}}}{attr}', font_value)
                            
                            # If direct font name is empty, extract from rFonts
                            if not target_run.font.name:
                                # Try to get font name from rFonts attributes, prioritize eastAsia
                                for attr in ['eastAsia', 'ascii', 'hAnsi']:
                                    font_name = source_rfonts.get(f'{{{w_ns}}}{attr}')
                                    if font_name:
                                        target_run.font.name = font_name
                                        break
                    
                    # If still no font name, try to get default font from source or target paragraph
                    if not target_run.font.name:
                        # Get font from source paragraph style
                        if (source_para and hasattr(source_para, 'style') and source_para.style and 
                            hasattr(source_para.style, 'font') and source_para.style.font and 
                            source_para.style.font.name):
                            target_run.font.name = source_para.style.font.name
                        # Get from target paragraph default font
                        elif (target_para and hasattr(target_para, '_default_font_name') and 
                              target_para._default_font_name):
                            target_run.font.name = target_para._default_font_name
                except Exception:
                    pass
                
                # Font size - improved handling for numeric font sizes
                try:
                    if hasattr(source_run.font, 'size') and source_run.font.size:
                        target_run.font.size = source_run.font.size
                    elif target_para and hasattr(target_para, '_default_font_size') and target_para._default_font_size:
                        # If run has no font size, use default size from paragraph style
                        target_run.font.size = target_para._default_font_size
                    else:
                        # Try to extract font size from XML if direct access fails
                        try:
                            if hasattr(source_run.font, '_element'):
                                w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                                sz_elem = source_run.font._element.find(f'.//{{{w_ns}}}sz')
                                if sz_elem is not None:
                                    sz_val = sz_elem.get(f'{{{w_ns}}}val')
                                    if sz_val:
                                        # Convert half-points to points (Word uses half-points internally)
                                        from docx.shared import Pt
                                        target_run.font.size = Pt(float(sz_val) / 2)
                        except Exception:
                            pass
                except Exception:
                    pass
                
                # Bold
                try:
                    if hasattr(source_run.font, 'bold') and source_run.font.bold is not None:
                        target_run.font.bold = bool(source_run.font.bold)
                except Exception:
                    pass
                
                # Italic
                try:
                    if hasattr(source_run.font, 'italic') and source_run.font.italic is not None:
                        target_run.font.italic = bool(source_run.font.italic)
                except Exception:
                    pass
                
                # Underline
                try:
                    if hasattr(source_run.font, 'underline') and source_run.font.underline is not None:
                        target_run.font.underline = source_run.font.underline
                except Exception:
                    pass
                
                # Font color
                try:
                    if hasattr(source_run.font, 'color') and source_run.font.color:
                        if hasattr(source_run.font.color, 'rgb') and source_run.font.color.rgb:
                            target_run.font.color.rgb = source_run.font.color.rgb
                        elif hasattr(source_run.font.color, 'theme_color') and source_run.font.color.theme_color:
                            target_run.font.color.theme_color = source_run.font.color.theme_color
                except Exception:
                    pass
                
                # Strikethrough
                try:
                    if hasattr(source_run.font, 'strike') and source_run.font.strike is not None:
                        target_run.font.strike = bool(source_run.font.strike)
                except Exception:
                    pass
                
                # Double strikethrough
                try:
                    if hasattr(source_run.font, 'double_strike') and source_run.font.double_strike is not None:
                        target_run.font.double_strike = bool(source_run.font.double_strike)
                except Exception:
                    pass
                
                # Superscript
                try:
                    if hasattr(source_run.font, 'superscript') and source_run.font.superscript is not None:
                        target_run.font.superscript = bool(source_run.font.superscript)
                except Exception:
                    pass
                
                # Subscript
                try:
                    if hasattr(source_run.font, 'subscript') and source_run.font.subscript is not None:
                        target_run.font.subscript = bool(source_run.font.subscript)
                except Exception:
                    pass
                
                # Small caps
                try:
                    if hasattr(source_run.font, 'small_caps') and source_run.font.small_caps is not None:
                        target_run.font.small_caps = bool(source_run.font.small_caps)
                except Exception:
                    pass
                
                # All caps
                try:
                    if hasattr(source_run.font, 'all_caps') and source_run.font.all_caps is not None:
                        target_run.font.all_caps = bool(source_run.font.all_caps)
                except Exception:
                    pass
                
                # Hidden text
                try:
                    if hasattr(source_run.font, 'hidden') and source_run.font.hidden is not None:
                        target_run.font.hidden = bool(source_run.font.hidden)
                except Exception:
                    pass
                
                # Character spacing
                try:
                    if hasattr(source_run.font, 'spacing') and source_run.font.spacing:
                        target_run.font.spacing = source_run.font.spacing
                except Exception:
                    pass
                
                # Character scaling
                try:
                    if hasattr(source_run.font, 'scale') and source_run.font.scale:
                        target_run.font.scale = source_run.font.scale
                except Exception:
                    pass
                
                # Highlight color
                try:
                    if hasattr(source_run.font, 'highlight_color') and source_run.font.highlight_color:
                        target_run.font.highlight_color = source_run.font.highlight_color
                except Exception:
                    pass
        except Exception as e:
            logger.debug(f"Failed to copy run format: {e}")
    
    def _copy_hyperlink(self, source_run, target_run, source_para, target_para):
        """Copy hyperlink from source run to target run"""
        try:
            # Check if source run is part of a hyperlink by examining the paragraph structure
            if hasattr(source_para, '_element') and source_para._element is not None:
                w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                r_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
                
                # Find all hyperlink elements in the paragraph
                hyperlink_elems = source_para._element.findall(f'.//{{{w_ns}}}hyperlink')
                
                for hyperlink_elem in hyperlink_elems:
                    # Get all text content from the hyperlink
                    hyperlink_text_elems = hyperlink_elem.findall(f'.//{{{w_ns}}}t')
                    hyperlink_texts = [t.text for t in hyperlink_text_elems if t.text]
                    hyperlink_full_text = ''.join(hyperlink_texts)
                    
                    # Check if this hyperlink text matches our source run text
                    if hyperlink_full_text == source_run.text:
                        # Found our run in a hyperlink, extract URL
                        r_id = hyperlink_elem.get(f'{{{r_ns}}}id')
                        anchor = hyperlink_elem.get(f'{{{w_ns}}}anchor')
                        
                        url = None
                        if r_id:
                            # Get the actual URL from document relationships
                            try:
                                source_doc_part = source_para._parent
                                if hasattr(source_doc_part, 'part') and hasattr(source_doc_part.part, 'rels'):
                                    rel = source_doc_part.part.rels[r_id]
                                    if hasattr(rel, 'target_ref'):
                                        url = rel.target_ref
                            except Exception as e:
                                logger.debug(f"Failed to get hyperlink URL from relationship: {e}")
                        elif anchor:
                            url = f"#{anchor}"
                        
                        if url:
                            # Add hyperlink to target document
                            self._add_hyperlink_to_run(target_run, url, target_run.text)
                            return
                        
        except Exception as e:
            logger.debug(f"Failed to copy hyperlink: {e}")
    
    def _add_hyperlink_to_run(self, run, url, text):
        """Add hyperlink to a run"""
        try:
            # Get the paragraph that contains this run
            paragraph = run._parent
            document = paragraph._parent
            
            # Ensure Hyperlink style exists
            self._ensure_hyperlink_style(document)
            
            # Find the exact run object in the paragraph's XML element
            run_index = None
            para_runs = paragraph._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            
            # Find the index of the run element in the paragraph's XML
            for i, run_elem in enumerate(para_runs):
                # Check if this is the same run element
                if run._element == run_elem:
                    run_index = i
                    break
            
            if run_index is not None:
                # Store run formatting
                font_name = run.font.name
                font_size = run.font.size
                font_bold = run.font.bold
                font_italic = run.font.italic
                font_underline = run.font.underline
                font_color = None
                if hasattr(run.font, 'color') and run.font.color:
                    if hasattr(run.font.color, 'rgb'):
                        font_color = run.font.color.rgb
                
                # Store the position where we need to insert the hyperlink
                # We need to find the position in the paragraph's children, not just runs
                para_children = list(paragraph._element)
                insert_position = None
                
                for i, child in enumerate(para_children):
                    if child == run._element:
                        insert_position = i
                        break
                
                # Remove the original run
                paragraph._element.remove(run._element)
                
                # Add hyperlink
                from docx.oxml.shared import qn
                from docx.oxml import parse_xml
                
                # Create hyperlink XML
                w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                r_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
                
                # Add relationship for external URL
                if url.startswith('http') or url.startswith('www.'):
                    # Add relationship to document
                    doc_part = paragraph._parent.part
                    rel_id = doc_part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
                    
                    hyperlink_xml = f'''
                    <w:hyperlink r:id="{rel_id}" xmlns:w="{w_ns}" xmlns:r="{r_ns}">
                        <w:r>
                            <w:rPr>
                                <w:rStyle w:val="Hyperlink"/>
                            </w:rPr>
                            <w:t>{text}</w:t>
                        </w:r>
                    </w:hyperlink>
                    '''
                else:
                    # Internal link (anchor)
                    anchor = url.lstrip('#')
                    hyperlink_xml = f'''
                    <w:hyperlink w:anchor="{anchor}" xmlns:w="{w_ns}">
                        <w:r>
                            <w:rPr>
                                <w:rStyle w:val="Hyperlink"/>
                            </w:rPr>
                            <w:t>{text}</w:t>
                        </w:r>
                    </w:hyperlink>
                    '''
                
                hyperlink_elem = parse_xml(hyperlink_xml)
                
                # Insert hyperlink at the correct position
                if insert_position is not None and insert_position < len(paragraph._element):
                    paragraph._element.insert(insert_position, hyperlink_elem)
                else:
                    paragraph._element.append(hyperlink_elem)
                
                # Apply original formatting to hyperlink run
                hyperlink_run = hyperlink_elem.find(f'.//{{{w_ns}}}r')
                if hyperlink_run is not None:
                    rpr = hyperlink_run.find(f'{{{w_ns}}}rPr')
                    if rpr is not None:
                        # Add font formatting
                        if font_name:
                            rfonts_xml = f'<w:rFonts w:ascii="{font_name}" w:eastAsia="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}" xmlns:w="{w_ns}"/>'
                            rfonts_elem = parse_xml(rfonts_xml)
                            rpr.append(rfonts_elem)
                        
                        if font_size:
                            # Convert points to half-points
                            half_points = int(font_size.pt * 2)
                            sz_xml = f'<w:sz w:val="{half_points}" xmlns:w="{w_ns}"/>'
                            sz_elem = parse_xml(sz_xml)
                            rpr.append(sz_elem)
                        
                        if font_bold:
                            b_xml = f'<w:b xmlns:w="{w_ns}"/>'
                            b_elem = parse_xml(b_xml)
                            rpr.append(b_elem)
                        
                        if font_italic:
                            i_xml = f'<w:i xmlns:w="{w_ns}"/>'
                            i_elem = parse_xml(i_xml)
                            rpr.append(i_elem)
                        
                        if font_underline:
                            u_xml = f'<w:u w:val="single" xmlns:w="{w_ns}"/>'
                            u_elem = parse_xml(u_xml)
                            rpr.append(u_elem)
                        
                        if font_color:
                            color_xml = f'<w:color w:val="{str(font_color)}" xmlns:w="{w_ns}"/>'
                            color_elem = parse_xml(color_xml)
                            rpr.append(color_elem)
                            
        except Exception as e:
            logger.warning(f"Failed to add hyperlink to run: {e}")
            # If hyperlink creation fails, at least preserve the text
            pass
    
    def _ensure_hyperlink_style(self, document):
        """Ensure Hyperlink style exists in the document"""
        try:
            # Check if Hyperlink style already exists
            for style in document.styles:
                if style.name == 'Hyperlink':
                    return  # Style already exists
            
            # Create Hyperlink style
            from docx.enum.style import WD_STYLE_TYPE
            from docx.shared import RGBColor
            
            hyperlink_style = document.styles.add_style('Hyperlink', WD_STYLE_TYPE.CHARACTER)
            hyperlink_style.font.color.rgb = RGBColor(0, 0, 255)  # Blue color
            hyperlink_style.font.underline = True
            
            logger.debug("Created Hyperlink style for document")
            
        except Exception as e:
            logger.warning(f"Failed to create Hyperlink style: {e}")
    
    def _get_images_from_run(self, run, document_part):
        """Extract image data from run"""
        images = []
        try:
            import xml.etree.ElementTree as ET
            from io import BytesIO
            
            # Parse run XML to find images
            run_xml = run._element.xml
            root = ET.fromstring(run_xml)
            namespace = {
                'a': "http://schemas.openxmlformats.org/drawingml/2006/main",
                'r': "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
                'wp': "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
            }
            
            # Find inline images
            inlines = root.findall('.//wp:inline', namespace)
            
            for inline in inlines:
                try:
                    # Get image embed ID
                    blips = inline.findall('.//a:blip', namespace)
                    for blip in blips:
                        embed_id = blip.get(f"{{{namespace['r']}}}embed")
                        if embed_id:
                            try:
                                if hasattr(document_part, 'rels') and embed_id in document_part.rels:
                                    image_part = document_part.rels[embed_id].target_part
                                else:
                                    continue
                            except (TypeError, AttributeError):
                                # Handle unhashable type error
                                continue
                            image_data = image_part.blob
                            
                            # Get image dimensions
                            extent = inline.find('.//wp:extent', namespace)
                            width = None
                            height = None
                            if extent is not None:
                                cx = extent.get('cx')
                                cy = extent.get('cy')
                                if cx and cy:
                                    # Convert EMU to Inches (1 inch = 914400 EMU)
                                     width = Inches(int(cx) / 914400)
                                     height = Inches(int(cy) / 914400)
                            
                            # Return data to caller (using BytesIO)
                            images.append({
                                'stream': BytesIO(image_data),
                                'width': width,
                                'height': height
                            })
                            logger.debug(f"Successfully extracted image: {embed_id}")
                                
                except Exception as e:
                    logger.warning(f"Failed to extract image data: {e}")
                    continue
                    
        except Exception as e:
            logger.warning(f"Failed to parse images in run: {e}")
            
        return images
    
    # ========================================================================
    # Table processing methods
    # ========================================================================
    
    def _copy_tables_in_range(self, source_doc: Document, target_doc: Document, paragraph_range: List[int]):
        """Copy tables within specified range, including images in tables"""
        try:
            # Get the order of all elements in the document (paragraphs and tables)
            document_elements = []
            
            # Traverse the document's body elements, collecting paragraphs and tables in order
            from docx.oxml.ns import qn
            body = source_doc._body._body
            
            para_index = 0
            table_index = 0
            
            for element in body:
                if element.tag == qn('w:p'):  # paragraph
                    document_elements.append(('paragraph', para_index))
                    para_index += 1
                elif element.tag == qn('w:tbl'):  # table
                    document_elements.append(('table', table_index))
                    table_index += 1
            
            # Find tables within the specified paragraph range
            tables_to_copy = []
            
            # Get the start and end paragraph indices of the chapter
            if not paragraph_range:
                return
                
            min_para = min(paragraph_range)
            max_para = max(paragraph_range)
            
            for i, (element_type, element_index) in enumerate(document_elements):
                if element_type == 'table':
                    # Check if the table is within the chapter's paragraph range
                    should_copy_table = False
                    
                    # Find the nearest paragraph before the table
                    prev_para_index = None
                    for j in range(i-1, -1, -1):
                        if document_elements[j][0] == 'paragraph':
                            prev_para_index = document_elements[j][1]
                            break
                    
                    # Find the nearest paragraph after the table
                    next_para_index = None
                    for j in range(i+1, len(document_elements)):
                        if document_elements[j][0] == 'paragraph':
                            next_para_index = document_elements[j][1]
                            break
                    
                    # Determine if the table belongs to the current chapter:
                    # 1. The paragraph before the table is within the current chapter range
                    # 2. The paragraph after the table is not within the current chapter range or the paragraph before the table is closer to the chapter start
                    if prev_para_index is not None and prev_para_index in paragraph_range:
                        # If there is no paragraph after the table, or the paragraph after is not in the current chapter, then the table belongs to the current chapter
                        if next_para_index is None or next_para_index not in paragraph_range:
                            should_copy_table = True
                        # If both paragraphs before and after the table are in the current chapter, then the table belongs to the current chapter
                        elif next_para_index in paragraph_range:
                            should_copy_table = True
                    
                    # If the table should be copied to the current chapter
                    if should_copy_table:
                        tables_to_copy.append(element_index)
            
            # Copy tables that meet the criteria
            for table_index in tables_to_copy:
                if table_index < len(source_doc.tables):
                    table = source_doc.tables[table_index]
                    new_table = target_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                    
                    # Copy table style
                    try:
                        if hasattr(table, 'style') and table.style:
                            style_name = table.style.name
                            logger.debug(f"Attempting to copy table style: {style_name}")
                            
                            # First ensure the target document has this style
                            target_style = None
                            for style in target_doc.styles:
                                if style.name == style_name:
                                    target_style = style
                                    break
                            
                            if target_style:
                                # Check if the existing style is complete
                                from docx.oxml.ns import qn
                                is_complete = False
                                if hasattr(target_style, '_element'):
                                    element = target_style._element
                                    has_tblPr = element.find(qn('w:tblPr')) is not None
                                    has_conditions = len(element.findall(qn('w:tblStylePr'))) > 0
                                    is_complete = has_tblPr or has_conditions
                                
                                if is_complete:
                                    # If the target document already has a complete style, use it directly
                                    try:
                                        new_table.style = target_style
                                        logger.debug(f"Successfully set table style: {style_name}")
                                    except Exception as e:
                                        logger.debug(f"Failed to set table style: {e}, using XML copy")
                                        self._copy_table_xml_style(table, new_table)
                                else:
                                    # If the style is incomplete, need to re-copy the style definition
                                    logger.debug(f"Target style {style_name} is incomplete, re-copying style definition")
                                    try:
                                        self._copy_table_style_definition(source_doc, target_doc, table.style)
                                        # Retry setting the style
                                        for style in target_doc.styles:
                                            if style.name == style_name:
                                                new_table.style = style
                                                logger.debug(f"Successfully set table style after re-copying style definition: {style_name}")
                                                break
                                        else:
                                            # If still fails, use XML copy
                                            logger.debug(f"Failed to re-copy style definition, using XML copy")
                                            self._copy_table_xml_style(table, new_table)
                                    except Exception as e:
                                        logger.debug(f"Failed to re-copy table style definition: {e}, using XML copy")
                                        self._copy_table_xml_style(table, new_table)
                            else:
                                # If the target document doesn't have this style, try to copy the style definition
                                try:
                                    self._copy_table_style_definition(source_doc, target_doc, table.style)
                                    # Retry setting the style
                                    for style in target_doc.styles:
                                        if style.name == style_name:
                                            new_table.style = style
                                            logger.debug(f"Successfully set table style after copying style definition: {style_name}")
                                            break
                                    else:
                                        # If still fails, use XML copy
                                        logger.debug(f"Style definition copy failed, using XML copy")
                                        self._copy_table_xml_style(table, new_table)
                                except Exception as e:
                                    logger.debug(f"Failed to copy table style definition: {e}, using XML copy")
                                    self._copy_table_xml_style(table, new_table)
                        else:
                            # If there's no style, still copy XML attributes to maintain formatting
                            self._copy_table_xml_style(table, new_table)
                        
                        # Copy table alignment
                        if hasattr(table, 'alignment'):
                            new_table.alignment = table.alignment
                    except Exception as e:
                        logger.warning(f"Table style copy failed: {e}")
                    
                    # Copy row and column dimensions
                    try:
                        for i, (source_row, target_row) in enumerate(zip(table.rows, new_table.rows)):
                            if hasattr(source_row, 'height'):
                                target_row.height = source_row.height
                    except Exception as e:
                        logger.warning(f"Table row height copy failed: {e}")
                    
                    try:
                        for j, (source_col, target_col) in enumerate(zip(table.columns, new_table.columns)):
                            if hasattr(source_col, 'width'):
                                target_col.width = source_col.width
                    except Exception as e:
                        logger.warning(f"Table column width copy failed: {e}")
                    
                    # Copy cell content and styles
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            target_cell = new_table.cell(i, j)
                            # Copy cell content, including images
                            self._copy_cell_content(cell, target_cell, source_doc.part)
                            # Copy cell style
                            self._copy_cell_style(cell, target_cell)
                    
                    logger.info(f"Copied table {table_index}, located within chapter range")
                        
        except Exception as e:
            logger.warning(f"Table copy warning: {e}")
    
    def _copy_table_xml_style(self, source_table, target_table):
        """Copy table's XML style attributes"""
        try:
            if hasattr(source_table, '_tbl') and hasattr(target_table, '_tbl'):
                source_tbl = source_table._tbl
                target_tbl = target_table._tbl
                
                # Copy table properties
                from docx.oxml.ns import qn
                source_tblPr = source_tbl.find(qn('w:tblPr'))
                if source_tblPr is not None:
                    # Get or create target table properties
                    target_tblPr = target_tbl.find(qn('w:tblPr'))
                    if target_tblPr is None:
                        from xml.etree.ElementTree import Element
                        target_tblPr = Element(qn('w:tblPr'))
                        target_tbl.insert(0, target_tblPr)
                    
                    # Copy table style
                    tblStyle = source_tblPr.find(qn('w:tblStyle'))
                    if tblStyle is not None:
                        # Remove existing style
                        existing_style = target_tblPr.find(qn('w:tblStyle'))
                        if existing_style is not None:
                            target_tblPr.remove(existing_style)
                        # Add new style
                        target_tblPr.append(copy.deepcopy(tblStyle))
                    
                    # Copy table width
                    tblW = source_tblPr.find(qn('w:tblW'))
                    if tblW is not None:
                        existing_tblW = target_tblPr.find(qn('w:tblW'))
                        if existing_tblW is not None:
                            target_tblPr.remove(existing_tblW)
                        target_tblPr.append(copy.deepcopy(tblW))
                    
                    # Copy table alignment
                    jc = source_tblPr.find(qn('w:jc'))
                    if jc is not None:
                        existing_jc = target_tblPr.find(qn('w:jc'))
                        if existing_jc is not None:
                            target_tblPr.remove(existing_jc)
                        target_tblPr.append(copy.deepcopy(jc))
                    
                    # Copy table borders
                    tblBorders = source_tblPr.find(qn('w:tblBorders'))
                    if tblBorders is not None:
                        existing_borders = target_tblPr.find(qn('w:tblBorders'))
                        if existing_borders is not None:
                            target_tblPr.remove(existing_borders)
                        target_tblPr.append(copy.deepcopy(tblBorders))
                    
                    # Copy table background
                    shd = source_tblPr.find(qn('w:shd'))
                    if shd is not None:
                        existing_shd = target_tblPr.find(qn('w:shd'))
                        if existing_shd is not None:
                            target_tblPr.remove(existing_shd)
                        target_tblPr.append(copy.deepcopy(shd))
                    
                    logger.debug("Table XML style copy completed")
                    
        except Exception as e:
            logger.warning(f"Table XML style copy failed: {e}")
    
    def _copy_table_style_definition(self, source_doc, target_doc, table_style):
        """Copy table style definition to target document"""
        try:
            style_name = table_style.name
            
            # Check if the target document already has this style and verify its completeness
            existing_style = None
            for style in target_doc.styles:
                if style.name == style_name:
                    existing_style = style
                    break
            
            if existing_style:
                # Check if the existing style is complete (has table properties or conditional formatting)
                from docx.oxml.ns import qn
                if hasattr(existing_style, '_element'):
                    element = existing_style._element
                    has_tblPr = element.find(qn('w:tblPr')) is not None
                    has_conditions = len(element.findall(qn('w:tblStylePr'))) > 0
                    
                    if has_tblPr or has_conditions:
                        logger.debug(f"Table style {style_name} already exists and is complete")
                        return
                    else:
                        logger.debug(f"Table style {style_name} exists but is incomplete, needs re-copying")
                        # Delete existing incomplete style
                        try:
                            target_doc.styles._element.remove(existing_style._element)
                            logger.debug(f"Deleted incomplete style: {style_name}")
                        except Exception as e:
                            logger.debug(f"Failed to delete incomplete style: {e}")
                else:
                    logger.debug(f"Table style {style_name} already exists")
                    return
            
            # Try to add new table style
            try:
                from docx.enum.style import WD_STYLE_TYPE
                new_style = target_doc.styles.add_style(style_name, WD_STYLE_TYPE.TABLE)
                logger.debug(f"Successfully added table style: {style_name}")
                
                # Copy basic style properties
                if hasattr(table_style, '_element') and hasattr(new_style, '_element'):
                    # Copy style's XML definition
                    from docx.oxml.ns import qn
                    
                    source_element = table_style._element
                    target_element = new_style._element
                    
                    # Copy table basic properties (tblPr)
                    tblPr = source_element.find(qn('w:tblPr'))
                    if tblPr is not None:
                        target_element.append(copy.deepcopy(tblPr))
                        logger.debug(f"Copied table basic properties (tblPr)")
                    
                    # Copy all conditional formatting (tblStylePr)
                    cond_formats = source_element.findall(qn('w:tblStylePr'))
                    for cond_format in cond_formats:
                        target_element.append(copy.deepcopy(cond_format))
                        cond_type = cond_format.get(qn('w:type'))
                        logger.debug(f"Copied conditional format: {cond_type}")
                    
                    # Copy other style properties
                    for child in source_element:
                        tag_name = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                        if tag_name not in ['tblPr', 'tblStylePr']:
                            # Check if target element already has this property
                            existing = target_element.find(child.tag)
                            if existing is None:
                                target_element.append(copy.deepcopy(child))
                                logger.debug(f"Copied style property: {tag_name}")
                        
                    logger.debug(f"Table style XML definition copy completed: {style_name}")
                    
            except Exception as e:
                logger.debug(f"Failed to add table style: {e}")
                # If unable to add style, try to copy style XML directly to document
                self._copy_style_xml_to_document(source_doc, target_doc, table_style)
                
        except Exception as e:
            logger.warning(f"Failed to copy table style definition: {e}")
    
    def _copy_style_xml_to_document(self, source_doc, target_doc, style):
        """Copy style XML directly to target document"""
        try:
            # Get style parts of source and target documents
            source_styles_part = source_doc.part.part_related_by('http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles')
            target_styles_part = target_doc.part.part_related_by('http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles')
            
            if source_styles_part and target_styles_part:
                from docx.oxml.ns import qn
                
                # Find XML definition of source style
                source_styles_root = source_styles_part.element
                target_styles_root = target_styles_part.element
                
                # Find specific style using style ID or name matching
                source_style_element = None
                for style_element in source_styles_root.findall(qn('w:style')):
                    style_id = style_element.get(qn('w:styleId'))
                    style_name_attr = style_element.find(qn('w:name'))
                    style_name = style_name_attr.get(qn('w:val')) if style_name_attr is not None else None
                    
                    # Match style ID or name
                    if (style_id and style_id == style.style_id) or \
                       (style_name and style_name == style.name) or \
                       (style_id and style_id == '11' and style.name == 'Grid Table 1 Light'):
                        source_style_element = style_element
                        logger.debug(f"Found source style: ID={style_id}, Name={style_name}")
                        break
                
                if source_style_element is not None:
                    # Find corresponding style element in target document
                    target_style_element = None
                    for existing_style in target_styles_root.findall(qn('w:style')):
                        existing_id = existing_style.get(qn('w:styleId'))
                        existing_name_attr = existing_style.find(qn('w:name'))
                        existing_name = existing_name_attr.get(qn('w:val')) if existing_name_attr is not None else None
                        
                        if (existing_id and existing_id == 'GridTable1Light') or \
                           (existing_name and existing_name == 'Grid Table 1 Light'):
                            target_style_element = existing_style
                            logger.debug(f"Found target style: ID={existing_id}, Name={existing_name}")
                            break
                    
                    if target_style_element is not None:
                        # Clear target style content, then copy all child elements from source style
                        for child in list(target_style_element):
                            if child.tag.split('}')[-1] not in ['name', 'styleId', 'type']:
                                target_style_element.remove(child)
                        
                        # Copy all attributes and child elements from source style (except basic identification info)
                        for child in source_style_element:
                            tag_name = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                            if tag_name not in ['name', 'styleId', 'type']:
                                target_style_element.append(copy.deepcopy(child))
                                logger.debug(f"Copied style element: {tag_name}")
                        
                        logger.debug(f"Successfully updated target style definition: {style.name}")
                    else:
                        logger.debug(f"Target style element not found: {style.name}")
                else:
                    logger.debug(f"Source style element not found: {style.name}")
                            
        except Exception as e:
            logger.debug(f"Direct style XML copy failed: {e}")
            import traceback
            logger.debug(traceback.format_exc())
    
    def _copy_cell_content(self, source_cell, target_cell, source_doc_part):
        """Copy table cell content, including text and images"""
        try:
            # Clear target cell
            target_cell.text = ""
            
            # Copy each paragraph
            for para in source_cell.paragraphs:
                # If it's the first paragraph, use existing paragraph
                if para == source_cell.paragraphs[0] and len(target_cell.paragraphs) > 0:
                    target_para = target_cell.paragraphs[0]
                    # Clear existing content
                    target_para.clear()
                else:
                    target_para = target_cell.add_paragraph()
                
                # Copy paragraph format
                if para.style:
                    try:
                        target_para.style = para.style
                    except Exception:
                        pass
                target_para.alignment = para.alignment
                
                # Copy runs
                for run in para.runs:
                    # Check if run contains images
                    images_in_run = self._get_images_from_run(run, source_doc_part)
                    
                    if images_in_run:
                        # If run contains images, add text first, then add images
                        if run.text:
                            new_run = target_para.add_run(run.text)
                            self._copy_run_format(run, new_run, para, target_para)
                        
                        # Add images
                        for image_data in images_in_run:
                            try:
                                new_run = target_para.add_run()
                                new_run.add_picture(image_data['stream'], width=image_data['width'], height=image_data['height'])
                                logger.info(f"Successfully copied image in table")
                            except Exception as e:
                                logger.warning(f"Table image copy failed: {e}")
                                # If image copy fails, add placeholder text
                                new_run = target_para.add_run("[Image]")
                    else:
                        # Regular text run
                        new_run = target_para.add_run(run.text)
                        self._copy_run_format(run, new_run, para, target_para)
                        
        except Exception as e:
            logger.warning(f"Cell content copy warning: {e}")
    
    def _copy_cell_style(self, source_cell, target_cell):
        """Copy table cell style"""
        try:
            from docx.oxml.ns import qn
            
            # Copy cell background color
            if hasattr(source_cell, '_tc') and hasattr(target_cell, '_tc'):
                source_tc = source_cell._tc
                target_tc = target_cell._tc
                
                # Copy cell properties
                tcPr = source_tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    # Create target cell properties
                    target_tcPr = target_tc.find(qn('w:tcPr'))
                    if target_tcPr is None:
                        target_tcPr = target_tc.insert(0, copy.deepcopy(tcPr))
                    else:
                        # Copy background color
                        shd = tcPr.find(qn('w:shd'))
                        if shd is not None:
                            target_shd = target_tcPr.find(qn('w:shd'))
                            if target_shd is not None:
                                target_tcPr.remove(target_shd)
                            target_tcPr.append(copy.deepcopy(shd))
                        
                        # Copy borders
                        tcBorders = tcPr.find(qn('w:tcBorders'))
                        if tcBorders is not None:
                            target_tcBorders = target_tcPr.find(qn('w:tcBorders'))
                            if target_tcBorders is not None:
                                target_tcPr.remove(target_tcBorders)
                            target_tcPr.append(copy.deepcopy(tcBorders))
                        
                        # Copy vertical alignment
                        vAlign = tcPr.find(qn('w:vAlign'))
                        if vAlign is not None:
                            target_vAlign = target_tcPr.find(qn('w:vAlign'))
                            if target_vAlign is not None:
                                target_tcPr.remove(target_vAlign)
                            target_tcPr.append(copy.deepcopy(vAlign))
                        
                        # Copy cell width
                        tcW = tcPr.find(qn('w:tcW'))
                        if tcW is not None:
                            target_tcW = target_tcPr.find(qn('w:tcW'))
                            if target_tcW is not None:
                                target_tcPr.remove(target_tcW)
                            target_tcPr.append(copy.deepcopy(tcW))
                            
        except Exception as e:
            logger.warning(f"Cell style copy warning: {e}")
    
    # ========================================================================
    # Utility methods
    # ========================================================================
    
    def _sanitize_filename(self, filename: str) -> str:
        """Clean filename, remove or replace invalid characters"""
        
        # Remove or replace various whitespace and special characters
        filename = filename.replace('\u3000', ' ')  # Full-width space
        filename = filename.replace('\t', ' ')      # Tab
        filename = filename.replace('\n', ' ')      # Newline
        filename = filename.replace('\r', ' ')      # Carriage return
        
        # Remove invalid characters in Windows filenames
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        
        # Remove control characters
        filename = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', filename)
        
        # Compress multiple spaces to single space
        filename = re.sub(r'\s+', ' ', filename)
        
        # Remove leading and trailing spaces
        filename = filename.strip()
        
        # If filename is empty, use default name
        if not filename:
            filename = '未命名章节'
        
        # Limit filename length (Windows path limitation)
        return filename[:100]

# ============================================================================
# Document processor class
# ============================================================================

class DocumentProcessor:
    """
    Document processor (performance optimized version)
    Responsible for batch processing multiple Word documents, using multi-threading to improve processing efficiency.
    Automatically optimizes thread count, supports error handling and progress tracking.
    """
    
    def __init__(self, input_dir: str, output_dir: str, 
                 file_thread_count: int = 4, chapter_thread_count: int = 2, min_level: int = 5):
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        
        # 根据CPU核心数优化线程数
        import multiprocessing
        cpu_count = multiprocessing.cpu_count()
        self.file_thread_count = min(file_thread_count, cpu_count)
        self.chapter_thread_count = min(chapter_thread_count, max(1, cpu_count // 2))
        
        self.splitter = WordDocumentSplitter(
            min_level=min_level,
            max_workers_docs=self.file_thread_count,
            max_workers_chapters=self.chapter_thread_count
        )
        
        # 创建输出目录
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"Processor configuration: document threads={self.file_thread_count}, chapter threads={self.chapter_thread_count}, CPU cores={cpu_count}")
    
    def process_all_documents(self):
        """
        Process all Word documents in the input directory
        Automatically scans the input directory, filters temporary files, and uses multi-threading for parallel processing.
        Processing results are logged to the log file.
        """
        # Get all Word documents
        all_files = list(self.input_dir.glob('*.docx')) + list(self.input_dir.glob('*.doc'))
        
        # Filter out temporary files starting with ~ (Word creates these when documents are open)
        doc_files = [f for f in all_files if not f.name.startswith('~')]
        
        if not doc_files:
            logger.warning(f"No Word documents found in {self.input_dir}")
            return
        
        logger.info(f"Found {len(doc_files)} documents, starting processing...")
        
        # Use thread pool to process multiple documents
        with ThreadPoolExecutor(max_workers=self.file_thread_count) as executor:
            futures = [executor.submit(self.process_single_document, doc_file) 
                      for doc_file in doc_files]
            
            for future in as_completed(futures):
                try:
                    result = future.result()
                    logger.info(f"Document processing completed: {result}")
                except Exception as e:
                    logger.error(f"Document processing failed: {e}")
    
    def process_single_document(self, doc_path: Path) -> str:
        """Process single document (performance optimized version)
        
        Args:
            doc_path: Document path
            
        Returns:
            str: Processing result description
            
        Note:
            Automatically creates document-specific output directory, uses multi-threading for chapter processing,
            includes error handling and resource cleanup mechanisms.
        """
        doc = None
        try:
            logger.info(f"Starting to process document: {doc_path}")
            
            # Analyze document structure
            doc, chapters = self.splitter.analyze_document_structure(str(doc_path))
            
            if not chapters:
                logger.warning(f"Document {doc_path} has no splittable chapters")
                return f"跳过: {doc_path.name}"
            
            # Create document-specific output directory
            doc_output_dir = self.output_dir / doc_path.stem
            doc_output_dir.mkdir(parents=True, exist_ok=True)
            
            # Dynamically adjust thread count (based on chapter count)
            optimal_workers = min(self.chapter_thread_count, len(chapters), 8)
            
            # Use thread pool to process chapters
            with ThreadPoolExecutor(max_workers=optimal_workers) as executor:
                futures = [executor.submit(self.splitter.create_chapter_document, 
                                         doc, chapter, str(doc_output_dir))
                          for chapter in chapters]
                
                created_files = []
                failed_count = 0
                
                for future in as_completed(futures):
                    try:
                        result = future.result(timeout=300)  # 5 minute timeout
                        created_files.append(result)
                    except Exception as e:
                        failed_count += 1
                        logger.error(f"Chapter processing failed: {e}")
                

                
                # Force garbage collection
                gc.collect()
            
            result_msg = f"完成: {doc_path.name} -> {len(created_files)} 个章节"
            if failed_count > 0:
                result_msg += f" (失败: {failed_count})"
            
            return result_msg
            
        except Exception as e:
            logger.error(f"Error occurred while processing document {doc_path}: {e}")
            return f"失败: {doc_path.name}"
        finally:
            if doc:
                del doc
            gc.collect()